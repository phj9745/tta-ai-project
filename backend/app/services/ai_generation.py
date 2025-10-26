from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Literal, Optional, Sequence
import asyncio
import csv
import base64
import io
import json
import logging
import mimetypes
import os
import re
import zipfile
from pathlib import Path
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, Iterable, List, Literal, Mapping
from xml.etree import ElementTree as ET

from fastapi import HTTPException, UploadFile
from docx import Document
from openai import (
    APIError,
    BadRequestError,
    OpenAI,
    OpenAIError,
    PermissionDeniedError,
    RateLimitError,
)

from ..config import Settings
from .excel_templates import TESTCASE_EXPECTED_HEADERS
from .excel_templates.utils import AI_CSV_DELIMITER
from .excel_templates.feature_list import normalize_feature_list_records
from .openai_payload import AttachmentMetadata, OpenAIMessageBuilder
from .prompt_config import (
    PromptBuiltinContext,
    PromptConfig,
    PromptConfigService,
    PromptResourcesConfig,
)
from .prompt_request_log import PromptRequestLogService


@dataclass
class BufferedUpload:
    name: str
    content: bytes
    content_type: str | None


@dataclass
class GeneratedCsv:
    filename: str
    content: bytes
    csv_text: str
    defect_summary: List["DefectSummaryEntry"] | None = None
    defect_images: Dict[int, List[BufferedUpload]] | None = None
    project_overview: str | None = None


TESTCASE_SCENARIO_SYSTEM_PROMPT = (
    "당신은 소프트웨어 QA 테스터입니다. 제공된 기능 설명과 참고 이미지를 바탕으로 "
    "실행 가능한 테스트 시나리오 후보를 정리합니다."
)

TESTCASE_FINALIZE_SYSTEM_PROMPT = (
    "당신은 소프트웨어 QA 테스터입니다. 기능별로 정리된 시나리오 요약을 바탕으로 "
    "테스트케이스 표를 완성합니다."
)


@dataclass
class UploadContext:
    upload: BufferedUpload
    metadata: Dict[str, Any] | None


@dataclass
class PromptContextPreview:
    descriptor: str
    doc_id: str | None
    include_in_attachment_list: bool
    metadata: Dict[str, Any]


@dataclass
class NormalizedDefect:
    index: int
    original_text: str
    polished_text: str


@dataclass(frozen=True)
class DefectSummaryAttachment:
    file_name: str
    original_file_name: str | None = None


@dataclass(frozen=True)
class DefectSummaryEntry:
    index: int
    original_text: str
    polished_text: str
    attachments: List[DefectSummaryAttachment]


@dataclass(frozen=True)
class DefectConversationTurn:
    role: Literal["user", "assistant"]
    text: str


@dataclass(frozen=True)
class DefectPromptResources:
    judgement_criteria: str | None = None
    output_example: str | None = None
    conversation: List[DefectConversationTurn] = field(default_factory=list)

logger = logging.getLogger(__name__)


class AIGenerationService:
    def __init__(
        self,
        settings: Settings,
        prompt_config_service: PromptConfigService | None = None,
        request_log_service: PromptRequestLogService | None = None,
    ):
        self._settings = settings
        if prompt_config_service is None:
            storage_path = settings.tokens_path.with_name("prompt_configs.json")
            prompt_config_service = PromptConfigService(storage_path)
        self._prompt_config_service = prompt_config_service
        self._client: OpenAI | None = None
        self._request_log_service = request_log_service

    def _get_client(self) -> OpenAI:
        if self._client is None:
            api_key = self._settings.openai_api_key
            if not api_key:
                raise HTTPException(status_code=500, detail="OpenAI API 키가 설정되어 있지 않습니다.")
            self._client = OpenAI(api_key=api_key)
        return self._client

    @staticmethod
    def _descriptor_from_context(
        context: UploadContext,
    ) -> tuple[str, str | None, bool, Dict[str, Any]]:
        metadata = context.metadata or {}
        role = str(metadata.get("role") or "").strip()
        label = str(
            metadata.get("label") or metadata.get("description") or ""
        ).strip()
        description = str(metadata.get("description") or "").strip()
        notes = str(metadata.get("notes") or "").strip()
        source_path = str(metadata.get("source_path") or "").strip()

        extension = AIGenerationService._extension(context.upload)

        if role == "additional":
            base_label = label or "추가 문서"
            descriptor = f"추가 문서: {base_label}"
        elif label:
            descriptor = label
        else:
            descriptor = context.upload.name

        if extension:
            descriptor = f"{descriptor} ({extension})"

        doc_id = (
            str(metadata.get("id")) if role == "required" and metadata.get("id") else None
        )

        include_in_attachment_list = bool(metadata.get("show_in_attachment_list", True))
        preview_metadata: Dict[str, Any] = {
            "label": label or context.upload.name,
            "description": description,
            "role": role,
            "extension": extension,
            "notes": notes,
            "source_path": source_path,
        }
        return descriptor, doc_id, include_in_attachment_list, preview_metadata

    @staticmethod
    def _extension(upload: BufferedUpload) -> str:
        extension = os.path.splitext(upload.name)[1].lstrip(".")
        if extension:
            extension = extension.upper()
        elif upload.content_type:
            subtype = upload.content_type.split("/")[-1]
            extension = subtype.upper()
        mapping = {"JPEG": "JPG"}
        return mapping.get(extension, extension)

    @staticmethod
    def _attachment_kind(upload: BufferedUpload) -> Literal["file", "image"]:
        content_type = (upload.content_type or "").split(";")[0].strip().lower()
        if content_type.startswith("image/"):
            return "image"

        extension = os.path.splitext(upload.name)[1].lower()
        if extension in {
            ".png",
            ".jpg",
            ".jpeg",
            ".gif",
            ".bmp",
            ".webp",
            ".tiff",
            ".tif",
            ".heic",
        }:
            return "image"

        return "file"

    @classmethod
    def _normalize_upload_for_openai(cls, upload: BufferedUpload) -> BufferedUpload:
        content_type = (upload.content_type or "").split(";")[0].strip().lower()
        extension = Path(upload.name).suffix.lower()
        if extension in {".html", ".htm"} or content_type == "text/html":
            return cls._html_to_pdf(upload)
        return upload

    @staticmethod
    def _html_to_pdf(upload: BufferedUpload) -> BufferedUpload:
        try:
            soup = BeautifulSoup(upload.content, "html.parser")
            text = soup.get_text(separator="\n", strip=True)
        except Exception:
            text = ""

        if not text:
            try:
                text = upload.content.decode("utf-8", errors="ignore")
            except Exception:
                text = ""

        lines: List[str] = []
        if text:
            previous_blank = False
            for raw_line in text.splitlines():
                stripped = raw_line.strip()
                if stripped:
                    lines.append(stripped)
                    previous_blank = False
                elif not previous_blank:
                    lines.append("")
                    previous_blank = True
        if not lines:
            lines = ["원본 HTML에서 텍스트를 추출하지 못했습니다."]

        pdf_bytes = AIGenerationService._lines_to_pdf(lines)

        stem = Path(upload.name).stem or "document"
        new_name = f"{stem}.pdf"

        return BufferedUpload(
            name=new_name,
            content=pdf_bytes,
            content_type="application/pdf",
        )

    @classmethod
    def _prepare_contexts_for_openai(
        cls, contexts: Iterable[UploadContext]
    ) -> List[UploadContext]:
        prepared: List[UploadContext] = []
        for context in contexts:
            normalized_upload = cls._normalize_upload_for_openai(context.upload)
            if normalized_upload is context.upload:
                prepared.append(context)
            else:
                metadata = dict(context.metadata) if context.metadata is not None else None
                prepared.append(UploadContext(upload=normalized_upload, metadata=metadata))
        return prepared

    @staticmethod
    def _context_summary(menu_id: str, contexts: List[PromptContextPreview]) -> str:
        if not contexts:
            return ""

        def describe(preferred_ids: List[str]) -> str:
            ordered: List[str] = []
            for doc_id in preferred_ids:
                match = next(
                    (
                        context.metadata.get("label")
                        or context.descriptor
                        for context in contexts
                        if context.doc_id == doc_id
                    ),
                    None,
                )
                if match:
                    ordered.append(match)
            if len(ordered) == len(preferred_ids):
                return ", ".join(ordered)
            return ", ".join(
                context.metadata.get("label") or context.descriptor
                for context in contexts
            )

        if menu_id == "feature-list":
            description = describe(["user-manual", "configuration", "vendor-feature-list"])
            return description

        if menu_id == "testcase-generation":
            description = describe(["user-manual", "configuration", "vendor-feature-list"])
            return description

        return ", ".join(
            context.metadata.get("label") or context.descriptor for context in contexts
        )

    @staticmethod
    def _build_context_previews(
        contexts: Iterable[UploadContext],
    ) -> List[PromptContextPreview]:
        previews: List[PromptContextPreview] = []
        for context in contexts:
            (
                descriptor,
                doc_id,
                include_in_attachment_list,
                metadata,
            ) = AIGenerationService._descriptor_from_context(context)
            cleaned = descriptor.strip() or context.upload.name
            previews.append(
                PromptContextPreview(
                    descriptor=cleaned,
                    doc_id=doc_id,
                    include_in_attachment_list=include_in_attachment_list,
                    metadata=metadata,
                )
            )
        return previews

    async def _upload_openai_file(self, client: OpenAI, context: UploadContext) -> str:
        upload = context.upload
        stream = io.BytesIO(upload.content)
        try:
            created = await asyncio.to_thread(
                client.files.create,
                file=(upload.name, stream),
                purpose="assistants",
            )
        except (APIError, OpenAIError) as exc:
            raise HTTPException(
                status_code=502,
                detail=f"OpenAI 파일 업로드 중 오류가 발생했습니다: {exc}",
            ) from exc
        except Exception as exc:  # pragma: no cover - 안전망
            logger.exception(
                "Unexpected error uploading file to OpenAI",
                extra={"file_name": upload.name},
            )
            raise HTTPException(
                status_code=502,
                detail="OpenAI 파일 업로드 중 예기치 않은 오류가 발생했습니다.",
            ) from exc

        file_id = getattr(created, "id", None)
        if not file_id and hasattr(created, "get"):
            try:
                file_id = created.get("id")  # type: ignore[call-arg]
            except Exception:  # pragma: no cover - dict-like guard
                file_id = None

        if not isinstance(file_id, str) or not file_id:
            raise HTTPException(
                status_code=502,
                detail="OpenAI 파일 업로드 응답에 file_id가 없습니다.",
            )

        return file_id

    async def _cleanup_openai_files(
        self, client: OpenAI, file_records: Iterable[tuple[str, bool]]
    ) -> None:
        for file_id, skip_cleanup in file_records:
            if skip_cleanup:
                continue
            try:
                await asyncio.to_thread(client.files.delete, file_id=file_id)
            except Exception as exc:  # pragma: no cover - 로그 목적
                logger.warning(
                    "Failed to delete temporary OpenAI file",
                    extra={"file_id": file_id, "error": str(exc)},
                )

    async def formalize_defect_notes(
        self,
        *,
        project_id: str,
        entries: List[Dict[str, str]],
        feature_context: str = "",
    ) -> List[NormalizedDefect]:
        if not entries:
            raise HTTPException(status_code=422, detail="정제할 결함 항목이 없습니다.")

        client = self._get_client()
        system_prompt = (
            "당신은 소프트웨어 시험 결과를 정리하는 품질 보증 문서 작성자입니다. "
            "사용자가 제공한 비격식 표현을 공문서에 적합한 격식 있는 문장으로 다듬어야 합니다."
        )

        bullet_lines: List[str] = []
        for entry in entries:
            index_value = entry.get("index")
            text_value = (entry.get("text") or "").strip()
            if not text_value:
                continue
            bullet_lines.append(f"{index_value}. {text_value}")

        if not bullet_lines:
            raise HTTPException(status_code=422, detail="결함 항목에서 내용을 찾을 수 없습니다.")

        context_prompt = ""
        stripped_context = feature_context.strip()
        if stripped_context:
            context_prompt = (
                "프로그램의 기능리스트 요약입니다. 결함을 다듬을 때 해당 기능의 목적과 범위를 고려하세요.\n"
                f"{stripped_context}\n\n"
            )

        base_prompt = (
            "다음 결함 설명을 공문서에 맞는 문장으로 다듬어 주세요.\n"
            "- 결과는 입력 순서를 유지한 번호 매기기 형식으로 작성하세요.\n"
            "- 각 줄은 '번호. 정제된 문장' 형태여야 합니다.\n"
            "- 존댓말 어미를 사용하고 한 문장 또는 한 문단으로 간결하게 정리하세요.\n"
            "- 번호 목록 이외의 설명이나 부가 문장은 작성하지 마세요.\n\n"
        )
        user_prompt = (
            base_prompt
            + (context_prompt or "")
            + "입력 결함 목록:\n"
            + "\n".join(bullet_lines)
        )

        messages = [
            OpenAIMessageBuilder.text_message("system", system_prompt),
            OpenAIMessageBuilder.text_message("user", user_prompt),
        ]

        try:
            response = await asyncio.to_thread(
                client.responses.create,
                model=self._settings.openai_model,
                input=messages,
                temperature=0.2,
                top_p=0.9,
                max_output_tokens=600,
            )
        except RateLimitError as exc:
            detail = self._format_openai_error(exc)
            raise HTTPException(
                status_code=429,
                detail=(
                    "OpenAI 사용량 한도를 초과했습니다. "
                    "관리자에게 문의하거나 잠시 후 다시 시도해 주세요."
                    f" ({detail})"
                ),
            ) from exc
        except (PermissionDeniedError, BadRequestError, APIError, OpenAIError) as exc:
            detail = self._format_openai_error(exc)
            raise HTTPException(
                status_code=502,
                detail=f"OpenAI 호출 중 오류가 발생했습니다: {detail}",
            ) from exc
        except Exception as exc:  # pragma: no cover - 안전망
            logger.exception(
                "Unexpected error while requesting OpenAI response",
                extra={"project_id": project_id, "menu_id": "defect-report-formalize"},
            )
            message = str(exc).strip()
            detail = (
                "OpenAI 응답을 가져오는 중 예기치 않은 오류가 발생했습니다."
                if not message
                else f"OpenAI 응답을 가져오는 중 예기치 않은 오류가 발생했습니다: {message}"
            )
            raise HTTPException(status_code=502, detail=detail) from exc

        response_text = self._extract_response_text(response) or ""

        if self._request_log_service is not None:
            try:
                self._request_log_service.record_request(
                    project_id=project_id,
                    menu_id="defect-report-formalize",
                    system_prompt=system_prompt,
                    user_prompt=user_prompt,
                    context_summary="결함 목록 정제",
                    response_text=response_text,
                )
            except Exception:  # pragma: no cover - logging must not fail request
                logger.exception(
                    "Failed to record prompt request log",
                    extra={"project_id": project_id, "menu_id": "defect-report-formalize"},
                )

        if not response_text:
            raise HTTPException(status_code=502, detail="OpenAI 응답에서 번호 목록을 찾을 수 없습니다.")

        polished_by_index: Dict[int, str] = {}
        numbered_pattern = re.compile(
            r"(?:^|\n)\s*(\d+)\.(.*?)(?=(?:\n\s*\d+\.)|\Z)",
            re.S,
        )

        for match in numbered_pattern.finditer(response_text):
            index_str, body = match.groups()
            try:
                index_value = int(index_str)
            except ValueError:
                continue
            polished_value = " ".join(body.strip().split())
            if not polished_value:
                continue
            polished_by_index[index_value] = polished_value

        if not polished_by_index:
            fallback_lines = [
                line.strip()
                for line in response_text.splitlines()
                if line.strip()
            ]
            for offset, line in enumerate(fallback_lines, start=1):
                polished_by_index[offset] = line

        results: List[NormalizedDefect] = []
        for entry in entries:
            index_value = int(entry.get("index", 0))
            original_text = str(entry.get("text") or "").strip()
            polished_text = polished_by_index.get(index_value, original_text)
            if not polished_text:
                continue
            results.append(
                NormalizedDefect(
                    index=index_value,
                    original_text=original_text,
                    polished_text=polished_text,
                )
            )

        if not results:
            raise HTTPException(status_code=502, detail="정제된 결함 결과가 비어 있습니다.")

        return results

    async def rewrite_defect_report_cell(
        self,
        *,
        project_id: str,
        column_key: str,
        column_label: str | None,
        original_value: str | None,
        instructions: str,
        row_values: Mapping[str, str] | None = None,
    ) -> str:
        normalized_instructions = (instructions or "").strip()
        if not normalized_instructions:
            raise HTTPException(status_code=422, detail="변경 요청 내용을 입력해 주세요.")

        normalized_column = (column_key or "").strip()
        if not normalized_column:
            raise HTTPException(status_code=422, detail="수정할 열 정보를 확인할 수 없습니다.")

        display_label = (column_label or "").strip() or normalized_column
        current_value = (original_value or "").strip()

        context_lines: List[str] = []
        if row_values:
            for key, value in row_values.items():
                if not isinstance(key, str) or key.strip() == "":
                    continue
                if key == normalized_column:
                    continue
                value_text = "" if value is None else str(value)
                value_text = value_text.strip()
                if not value_text:
                    continue
                context_lines.append(f"- {key}: {value_text}")

        system_prompt = (
            "당신은 소프트웨어 시험 결과를 정리하는 결함 리포트 편집자입니다. "
            "각 항목은 명확하고 공문서에 적합한 어조를 유지해야 합니다."
        )

        prompt_parts: List[str] = []
        if context_lines:
            prompt_parts.append("행의 다른 항목:\n" + "\n".join(context_lines))
        prompt_parts.append(f"현재 '{display_label}' 값: {current_value or '없음'}")
        prompt_parts.append(f"사용자 요청: {normalized_instructions}")
        prompt_parts.append(
            "위 정보를 바탕으로 해당 셀에 들어갈 문장을 공문서 어조로 작성해 주세요.\n"
            "- 출력은 수정된 셀 내용만 제공하세요.\n"
            "- 필요 시 존댓말을 사용하고 문장은 간결하게 유지하세요."
        )

        user_prompt = "\n\n".join(prompt_parts)

        client = self._get_client()
        messages = [
            OpenAIMessageBuilder.text_message("system", system_prompt),
            OpenAIMessageBuilder.text_message("user", user_prompt),
        ]

        try:
            response = await asyncio.to_thread(
                client.responses.create,
                model=self._settings.openai_model,
                input=messages,
                temperature=0.2,
                top_p=0.9,
                max_output_tokens=400,
            )
        except RateLimitError as exc:
            detail = self._format_openai_error(exc)
            raise HTTPException(
                status_code=429,
                detail=(
                    "OpenAI 사용량 한도를 초과했습니다. "
                    "관리자에게 문의하거나 잠시 후 다시 시도해 주세요."
                    f" ({detail})"
                ),
            ) from exc
        except (PermissionDeniedError, BadRequestError, APIError, OpenAIError) as exc:
            detail = self._format_openai_error(exc)
            raise HTTPException(
                status_code=502,
                detail=f"OpenAI 호출 중 오류가 발생했습니다: {detail}",
            ) from exc
        except Exception as exc:  # pragma: no cover - 안전망
            logger.exception(
                "Unexpected error while requesting OpenAI response",
                extra={
                    "project_id": project_id,
                    "menu_id": "defect-report-rewrite",
                    "column": normalized_column,
                },
            )
            message = str(exc).strip()
            detail = (
                "OpenAI 응답을 가져오는 중 예기치 않은 오류가 발생했습니다."
                if not message
                else f"OpenAI 응답을 가져오는 중 예기치 않은 오류가 발생했습니다: {message}"
            )
            raise HTTPException(status_code=502, detail=detail) from exc

        response_text = self._extract_response_text(response) or ""

        if self._request_log_service is not None:
            try:
                self._request_log_service.record_request(
                    project_id=project_id,
                    menu_id="defect-report-rewrite",
                    system_prompt=system_prompt,
                    user_prompt=user_prompt,
                    context_summary=f"{display_label} 셀 수정",
                    response_text=response_text,
                )
            except Exception:  # pragma: no cover - logging must not fail request
                logger.exception(
                    "Failed to record prompt request log",
                    extra={"project_id": project_id, "menu_id": "defect-report-rewrite"},
                )

        updated_value = response_text.strip()
        if not updated_value:
            raise HTTPException(status_code=502, detail="OpenAI 응답에서 수정된 텍스트를 찾을 수 없습니다.")

        return updated_value

    @staticmethod
    def _sanitize_csv(text: str) -> str:
        cleaned = text.strip()
        fence_match = re.search(r"```(?:csv)?\s*(.*?)```", cleaned, re.DOTALL | re.IGNORECASE)
        if fence_match:
            cleaned = fence_match.group(1).strip()
        return cleaned

    @staticmethod
    def _sanitize_json(text: str) -> str:
        cleaned = text.strip()
        fence_match = re.search(r"```(?:json)?\s*(.*?)```", cleaned, re.DOTALL | re.IGNORECASE)
        if fence_match:
            cleaned = fence_match.group(1).strip()
        return cleaned

    async def suggest_testcase_scenarios(
        self,
        *,
        project_id: str,
        major_category: str,
        middle_category: str,
        minor_category: str,
        feature_description: str,
        project_overview: str,
        scenario_count: int,
        attachments: Sequence[UploadFile],
    ) -> List[Dict[str, str]]:
        normalized_count = max(1, min(5, scenario_count))
        buffered_uploads: List[BufferedUpload] = []
        metadata_entries: List[Dict[str, Any]] = []

        for index, upload in enumerate(attachments, start=1):
            try:
                content = await upload.read()
            finally:
                await upload.close()

            name = upload.filename or f"attachment-{index}"
            buffered_uploads.append(
                BufferedUpload(
                    name=name,
                    content=content,
                    content_type=upload.content_type,
                )
            )
            metadata_entries.append(
                {
                    "label": f"{minor_category or '소분류'} 참고 자료 {index}",
                    "role": "additional",
                }
            )

        contexts: List[UploadContext] = [
            UploadContext(upload=upload, metadata=metadata)
            for upload, metadata in zip(buffered_uploads, metadata_entries)
        ]

        client = self._get_client()
        uploaded_records: List[tuple[str, bool]] = []
        attachments_payload: List[AttachmentMetadata] = []

        try:
            for context in contexts:
                kind = self._attachment_kind(context.upload)
                if kind == "image":
                    attachments_payload.append(
                        {
                            "kind": "image",
                            "image_url": self._image_data_url(context.upload),
                        }
                    )
                else:
                    file_id = await self._upload_openai_file(client, context)
                    uploaded_records.append((file_id, False))
                    attachments_payload.append({"kind": kind, "file_id": file_id})

            feature_lines = [
                f"대분류: {major_category or '-'}",
                f"중분류: {middle_category or '-'}",
                f"소분류: {minor_category or '-'}",
            ]
            description_text = feature_description.strip() or "(기능 설명이 제공되지 않았습니다.)"
            overview_text = project_overview.strip() or "(프로젝트 개요가 제공되지 않았습니다.)"

            instructions = [
                f"다음 기능에 대해 {normalized_count}개의 테스트 시나리오 후보를 JSON으로 작성해 주세요.",
                "각 시나리오는 '테스트 시나리오', '입력(사전조건 포함)', '기대 출력(사후조건 포함)' 키를 포함한 객체여야 합니다.",
                "전체 응답은 {\"scenarios\": [...]} 형태의 JSON 한 개만 반환하고 JSON 외 텍스트는 추가하지 마세요.",
                "'테스트 시나리오' 값은 테스트 목적을 한 문장으로 명확하게 설명해야 합니다.",
                "'입력(사전조건 포함)' 값은 실제 예시 데이터를 포함한 단계 번호 목록을 '1. ...' 형식으로 작성하고 줄바꿈으로 구분하세요.",
                "'기대 출력(사후조건 포함)' 값은 기대되는 시스템 반응을 한 문장으로 요약하세요.",
                "중복되거나 의미가 겹치는 시나리오는 피하세요.",
            ]

            user_prompt_parts = [
                "프로젝트 개요:",
                overview_text,
                "",
                "기능 분류:",
                "\n".join(feature_lines),
                "",
                "기능 설명:",
                description_text,
                "",
                "지시사항:",
                "\n".join(f"- {line}" for line in instructions),
            ]

            user_prompt = "\n".join(part for part in user_prompt_parts if part is not None)

            messages = [
                OpenAIMessageBuilder.text_message("system", TESTCASE_SCENARIO_SYSTEM_PROMPT),
                OpenAIMessageBuilder.text_message(
                    "user",
                    user_prompt,
                    attachments=attachments_payload if attachments_payload else None,
                ),
            ]

            normalized_messages = OpenAIMessageBuilder.normalize_messages(messages)

            try:
                response = await asyncio.to_thread(
                    client.responses.create,
                    model=self._settings.openai_model,
                    input=normalized_messages,
                    temperature=0.2,
                    top_p=0.9,
                    max_output_tokens=800,
                )
            except RateLimitError as exc:
                detail = self._format_openai_error(exc)
                raise HTTPException(
                    status_code=429,
                    detail=(
                        "OpenAI 사용량 한도를 초과했습니다. "
                        "관리자에게 문의하거나 잠시 후 다시 시도해 주세요."
                        f" ({detail})"
                    ),
                ) from exc
            except (PermissionDeniedError, BadRequestError, APIError, OpenAIError) as exc:
                detail = self._format_openai_error(exc)
                raise HTTPException(
                    status_code=502,
                    detail=f"OpenAI 호출 중 오류가 발생했습니다: {detail}",
                ) from exc
            except Exception as exc:  # pragma: no cover - 안전망
                logger.exception(
                    "Unexpected error while requesting scenario suggestions",
                    extra={"project_id": project_id},
                )
                raise HTTPException(
                    status_code=502,
                    detail="테스트 시나리오를 생성하는 중 예기치 않은 오류가 발생했습니다.",
                ) from exc

            response_text = self._extract_response_text(response) or ""
            cleaned = self._sanitize_json(response_text)
            try:
                payload = json.loads(cleaned)
            except json.JSONDecodeError as exc:
                raise HTTPException(
                    status_code=502,
                    detail="OpenAI 응답을 JSON으로 해석하지 못했습니다.",
                ) from exc

            scenarios_raw: Any
            if isinstance(payload, dict):
                scenarios_raw = payload.get("scenarios")
            else:
                scenarios_raw = payload

            if not isinstance(scenarios_raw, Sequence):
                raise HTTPException(
                    status_code=502,
                    detail="OpenAI 응답에서 시나리오 목록을 찾을 수 없습니다.",
                )

            normalized: List[Dict[str, str]] = []
            for entry in scenarios_raw:
                if not isinstance(entry, Mapping):
                    continue
                scenario_text = str(
                    entry.get("테스트 시나리오")
                    or entry.get("scenario")
                    or ""
                ).strip()
                input_text = str(
                    entry.get("입력(사전조건 포함)")
                    or entry.get("input")
                    or ""
                ).strip()
                expected_text = str(
                    entry.get("기대 출력(사후조건 포함)")
                    or entry.get("expected")
                    or ""
                ).strip()
                if not scenario_text:
                    continue
                normalized.append(
                    {
                        "scenario": scenario_text,
                        "input": input_text,
                        "expected": expected_text,
                    }
                )

            if not normalized:
                raise HTTPException(
                    status_code=502,
                    detail="OpenAI 응답에서 유효한 테스트 시나리오를 찾을 수 없습니다.",
                )

            return normalized
        finally:
            if uploaded_records:
                await self._cleanup_openai_files(client, uploaded_records)

    async def rewrite_testcase_scenarios(
        self,
        *,
        project_id: str,
        project_overview: str,
        major_category: str,
        middle_category: str,
        minor_category: str,
        feature_description: str,
        scenarios: Sequence[Mapping[str, Any]],
        instructions: str,
        conversation: Sequence[Mapping[str, str]] | None = None,
    ) -> Dict[str, Any]:
        normalized_instructions = (instructions or "").strip()
        if not normalized_instructions:
            raise HTTPException(status_code=422, detail="변경 요청 내용을 입력해 주세요.")

        normalized_scenarios: List[Dict[str, str]] = []
        for entry in scenarios:
            if not isinstance(entry, Mapping):
                continue
            scenario_text = str(
                entry.get("scenario")
                or entry.get("테스트 시나리오")
                or ""
            ).strip()
            input_text = str(
                entry.get("input")
                or entry.get("입력(사전조건 포함)")
                or ""
            ).strip()
            expected_text = str(
                entry.get("expected")
                or entry.get("기대 출력(사후조건 포함)")
                or ""
            ).strip()
            if not scenario_text:
                continue
            normalized_scenarios.append(
                {
                    "scenario": scenario_text,
                    "input": input_text,
                    "expected": expected_text,
                }
            )

        if not normalized_scenarios:
            raise HTTPException(status_code=422, detail="수정할 테스트케이스가 없습니다.")

        scenario_lines: List[str] = []
        for index, entry in enumerate(normalized_scenarios, start=1):
            scenario_lines.append(f"{index}. 테스트 시나리오: {entry['scenario']}")
            scenario_lines.append(
                "   입력(사전조건 포함): "
                + (entry["input"] or "-")
            )
            scenario_lines.append(
                "   기대 출력(사후조건 포함): "
                + (entry["expected"] or "-")
            )

        overview_text = project_overview.strip() or "(프로젝트 개요가 제공되지 않았습니다.)"
        feature_lines = [
            f"대분류: {major_category or '-'}",
            f"중분류: {middle_category or '-'}",
            f"소분류: {minor_category or '-'}",
        ]
        description_text = feature_description.strip() or "(기능 설명이 제공되지 않았습니다.)"

        response_instructions = [
            "아래 JSON 형식으로만 응답하세요.",
            (
                '{"reply": "요약 또는 변경 이유", "scenarios": '
                '[{"테스트 시나리오": "...", "입력(사전조건 포함)": "...", '
                '"기대 출력(사후조건 포함)": "..."}, ...]}'
            ),
            "scenarios 배열 길이는 최소 1개 이상이어야 하며 가능하면 현재 개수와 동일하게 유지하세요.",
            "각 항목은 한글 레이블을 그대로 사용하고 줄바꿈은 그대로 유지하세요.",
            "reply는 1~3문장으로 변경 사항을 요약하세요.",
        ]

        user_parts = [
            "프로젝트 개요:",
            overview_text,
            "",
            "기능 분류:",
            "\n".join(feature_lines),
            "",
            "기능 설명:",
            description_text,
            "",
            "현재 테스트케이스:",
            "\n".join(scenario_lines),
            "",
            "사용자 요청:",
            normalized_instructions,
            "",
            "응답 형식 지침:",
            "\n".join(f"- {line}" for line in response_instructions),
        ]

        user_prompt = "\n".join(part for part in user_parts if part is not None)

        system_prompt = (
            "당신은 소프트웨어 테스트 전문가입니다. "
            "사용자의 테스트케이스를 개선하고 명확하게 다듬어 주세요."
        )

        messages = [OpenAIMessageBuilder.text_message("system", system_prompt)]

        if conversation:
            for entry in conversation:
                role = entry.get("role")
                text = str(entry.get("text") or "").strip()
                if role not in {"user", "assistant"}:
                    continue
                if not text:
                    continue
                messages.append(OpenAIMessageBuilder.text_message(str(role), text))

        messages.append(OpenAIMessageBuilder.text_message("user", user_prompt))

        client = self._get_client()

        try:
            response = await asyncio.to_thread(
                client.responses.create,
                model=self._settings.openai_model,
                input=messages,
                temperature=0.2,
                top_p=0.9,
                max_output_tokens=900,
            )
        except RateLimitError as exc:
            detail = self._format_openai_error(exc)
            raise HTTPException(
                status_code=429,
                detail=(
                    "OpenAI 사용량 한도를 초과했습니다. "
                    "관리자에게 문의하거나 잠시 후 다시 시도해 주세요."
                    f" ({detail})"
                ),
            ) from exc
        except (PermissionDeniedError, BadRequestError, APIError, OpenAIError) as exc:
            detail = self._format_openai_error(exc)
            raise HTTPException(
                status_code=502,
                detail=f"OpenAI 호출 중 오류가 발생했습니다: {detail}",
            ) from exc
        except Exception as exc:  # pragma: no cover - 안전망
            logger.exception(
                "Unexpected error while requesting testcase rewrite",
                extra={
                    "project_id": project_id,
                    "menu_id": "testcase-generation",
                },
            )
            raise HTTPException(
                status_code=502,
                detail="테스트케이스를 다시 작성하는 중 예기치 않은 오류가 발생했습니다.",
            ) from exc

        response_text = self._extract_response_text(response) or ""
        cleaned = self._sanitize_json(response_text)

        try:
            payload = json.loads(cleaned)
        except json.JSONDecodeError as exc:
            raise HTTPException(
                status_code=502,
                detail="OpenAI 응답을 JSON으로 해석하지 못했습니다.",
            ) from exc

        reply_text = ""
        scenarios_payload: Any = []

        if isinstance(payload, Mapping):
            reply_text = str(
                payload.get("reply")
                or payload.get("message")
                or ""
            ).strip()
            scenarios_payload = payload.get("scenarios")
        else:
            scenarios_payload = payload

        if not isinstance(scenarios_payload, Sequence):
            raise HTTPException(
                status_code=502,
                detail="OpenAI 응답에서 수정된 테스트케이스를 찾을 수 없습니다.",
            )

        normalized_results: List[Dict[str, str]] = []
        for entry in scenarios_payload:
            if not isinstance(entry, Mapping):
                continue
            scenario_text = str(
                entry.get("테스트 시나리오")
                or entry.get("scenario")
                or ""
            ).strip()
            input_text = str(
                entry.get("입력(사전조건 포함)")
                or entry.get("input")
                or ""
            ).strip()
            expected_text = str(
                entry.get("기대 출력(사후조건 포함)")
                or entry.get("expected")
                or ""
            ).strip()
            if not scenario_text:
                continue
            normalized_results.append(
                {
                    "scenario": scenario_text,
                    "input": input_text,
                    "expected": expected_text,
                }
            )

        if not normalized_results:
            raise HTTPException(
                status_code=502,
                detail="OpenAI 응답에서 유효한 테스트케이스를 찾을 수 없습니다.",
            )

        return {
            "reply": reply_text,
            "scenarios": normalized_results,
        }

    async def generate_testcases_from_scenarios(
        self,
        *,
        project_id: str,
        project_overview: str,
        groups: Sequence[Mapping[str, Any]],
    ) -> GeneratedCsv:
        if not groups:
            raise HTTPException(status_code=422, detail="생성할 시나리오가 제공되지 않았습니다.")

        summary_blocks: List[str] = []
        total_scenarios = 0

        for group in groups:
            if not isinstance(group, Mapping):
                continue
            major = str(group.get("majorCategory") or "").strip()
            middle = str(group.get("middleCategory") or "").strip()
            minor = str(group.get("minorCategory") or "").strip()
            description = str(group.get("featureDescription") or "").strip()
            scenarios = group.get("scenarios")
            if not isinstance(scenarios, Sequence):
                continue

            scenario_lines: List[str] = []
            for index, entry in enumerate(scenarios, start=1):
                if not isinstance(entry, Mapping):
                    continue
                scenario_text = str(entry.get("scenario") or "").strip()
                input_text = str(entry.get("input") or "").strip()
                expected_text = str(entry.get("expected") or "").strip()
                if not scenario_text:
                    continue
                total_scenarios += 1
                scenario_lines.append(
                    "\n".join(
                        [
                            f"{index}. 시나리오: {scenario_text}",
                            f"   입력: {input_text or '-'}",
                            f"   기대 출력: {expected_text or '-'}",
                        ]
                    )
                )

            if not scenario_lines:
                continue

            block_parts = [
                f"대분류: {major or '-'}",
                f"중분류: {middle or '-'}",
                f"소분류: {minor or '-'}",
            ]
            if description:
                block_parts.append(f"기능 설명: {description}")
            block_parts.append("시나리오:")
            block_parts.extend(scenario_lines)
            summary_blocks.append("\n".join(block_parts))

        if not summary_blocks or total_scenarios == 0:
            raise HTTPException(status_code=422, detail="시나리오 요약이 비어 있습니다.")

        summary_text = "\n\n".join(summary_blocks)
        overview_text = project_overview.strip() or "(프로젝트 개요가 제공되지 않았습니다.)"
        headers_text = ", ".join(TESTCASE_EXPECTED_HEADERS)

        user_prompt_parts = [
            "프로젝트 개요:",
            overview_text,
            "",
            "기능별 시나리오 요약:",
            summary_text,
            "",
            "작성 지침:",
            "- 위 시나리오를 모두 포함하여 테스트케이스를 작성하세요.",
            f"- CSV 열은 {headers_text} 순서를 따릅니다.",
            "- 각 열은 파이프(|) 기호로 구분합니다.",
            "- 각 소분류 순서에 따라 테스트 케이스 ID 접두사를 TC-XXX-YYY 형식(예: TC-001-001)으로 부여하고",
            "  XXX는 소분류 그룹 번호(1부터 시작), YYY는 그룹 내 순번(1부터 시작)으로 3자리 숫자로 작성하세요.",
            "- '테스트 시나리오' 열은 '모든 입력필드에 유효한 값을 입력하여 기업이 정상적으로 생성되는지 확인'처럼",
            "  간결하고 자연스러운 한 문장으로 작성하세요.",
            "- '입력(사전조건 포함)' 열은 실제 예시값을 포함한 단계 번호 목록을 작성하고 각 단계는",
            "  '1. ...' 형식으로 시작하며 줄바꿈으로 구분하세요.",
            "- '기대 출력(사후조건 포함)' 열은 기대 결과를 한 문장으로 요약하고 안내 문구나 불필요한 설명을",
            "  추가하지 마세요.",
            "- 테스트 결과는 기본값으로 '미실행'을 사용하고 상세 테스트 결과와 비고는 비워 두세요.",
            "- 여러 줄이 필요한 열은 CSV 규칙에 맞게 큰따옴표로 감싸고 실제 줄바꿈 문자(엔터)를 사용하세요.",
            "- 출력 예시는 아래와 같이 작성합니다 (각 열은 파이프(|)로 구분됩니다):",
            "  TC-001-001 | 모든 입력필드에 유효한 값을 입력하여 기업이 정상적으로 생성되는지 확인 | \"1. 모든 입력필드에 유효한 값 입력\\n기업명: test\\n기업코드: TEST1\\n대표명: 홍길동\\n직급: 과장\\n주소: 서울특별시 마포구\\n연락처: 010-1234-5678\\n이메일: test1@gmail.com\\n팩스 번호: 02-123-4567\\n설명: 테스트\\n2. '생성' 버튼 클릭\" | 기업이 정상적으로 생성됨 | 미실행 |  | ",
            "- CSV 이외의 다른 텍스트나 설명을 포함하지 마세요.",
        ]

        user_prompt = "\n".join(user_prompt_parts)

        client = self._get_client()
        messages = [
            OpenAIMessageBuilder.text_message("system", TESTCASE_FINALIZE_SYSTEM_PROMPT),
            OpenAIMessageBuilder.text_message("user", user_prompt),
        ]

        normalized_messages = OpenAIMessageBuilder.normalize_messages(messages)

        try:
            response = await asyncio.to_thread(
                client.responses.create,
                model=self._settings.openai_model,
                input=normalized_messages,
                temperature=0.2,
                top_p=0.9,
                max_output_tokens=1800,
            )
        except RateLimitError as exc:
            detail = self._format_openai_error(exc)
            raise HTTPException(
                status_code=429,
                detail=(
                    "OpenAI 사용량 한도를 초과했습니다. "
                    "관리자에게 문의하거나 잠시 후 다시 시도해 주세요."
                    f" ({detail})"
                ),
            ) from exc
        except (PermissionDeniedError, BadRequestError, APIError, OpenAIError) as exc:
            detail = self._format_openai_error(exc)
            raise HTTPException(
                status_code=502,
                detail=f"OpenAI 호출 중 오류가 발생했습니다: {detail}",
            ) from exc
        except Exception as exc:  # pragma: no cover - 안전망
            logger.exception(
                "Unexpected error while finalising testcases",
                extra={"project_id": project_id},
            )
            raise HTTPException(
                status_code=502,
                detail="테스트케이스를 완성하는 중 예기치 않은 오류가 발생했습니다.",
            ) from exc

        response_text = self._extract_response_text(response) or ""
        sanitized = self._sanitize_csv(response_text)
        if not sanitized:
            raise HTTPException(status_code=502, detail="OpenAI 응답에서 CSV를 찾을 수 없습니다.")

        encoded = sanitized.encode("utf-8-sig")
        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        safe_project = re.sub(r"[^A-Za-z0-9_-]+", "_", project_id)
        filename = f"{safe_project}_testcase_workflow_{timestamp}.csv"

        return GeneratedCsv(
            filename=filename,
            content=encoded,
            csv_text=sanitized,
        )

    @staticmethod
    def _extract_feature_list_project_overview(
        csv_text: str,
    ) -> tuple[str, str | None]:
        project_overview: str | None = None
        rows_to_keep: list[list[str]] = []
        awaiting_overview_value = False

        stream = io.StringIO(csv_text)
        reader = csv.reader(stream, delimiter=AI_CSV_DELIMITER)

        for raw_row in reader:
            row = [cell.strip() for cell in raw_row]

            if not any(row):
                # Skip completely empty rows altogether.
                continue

            if awaiting_overview_value and project_overview is None:
                candidate_indexes = [idx for idx, cell in enumerate(row) if cell]
                if len(candidate_indexes) == 1:
                    candidate = row[candidate_indexes[0]].strip()
                    if candidate:
                        project_overview = candidate
                        awaiting_overview_value = False
                        continue
                awaiting_overview_value = False

            if project_overview is None and row:
                first_cell = row[0].lstrip("\ufeff").strip()
                colon_match = re.match(
                    r"^(?:프로젝트\s*)?개요\s*[:：\-]\s*(.+)$",
                    first_cell,
                    re.IGNORECASE,
                )
                if colon_match:
                    candidate = colon_match.group(1).strip()
                    if candidate:
                        project_overview = candidate
                        continue

                normalized_key = re.sub(r"\s+", "", first_cell.lower())
                if normalized_key in {"프로젝트개요", "개요"}:
                    remainder = next((cell for cell in row[1:] if cell), "").strip()
                    if remainder:
                        project_overview = remainder
                        continue
                    awaiting_overview_value = True
                    continue

            rows_to_keep.append(raw_row)

        if project_overview is None:
            fallback_match = re.search(
                r"(?:^|\n)\s*(?:프로젝트\s*)?개요\s*(?:[:：\-]\s*)?(.+)",
                csv_text,
            )
            if fallback_match:
                project_overview = fallback_match.group(1).strip()

        if not rows_to_keep:
            return "", project_overview

        output = io.StringIO()
        writer = csv.writer(output, lineterminator="\n", delimiter=AI_CSV_DELIMITER)
        for row in rows_to_keep:
            writer.writerow(row)
        return output.getvalue().strip(), project_overview

    @staticmethod
    def _format_feature_list_program_overview(
        records: Sequence[Mapping[str, str]],
        raw_overview: str | None,
        *,
        max_features: int = 6,
    ) -> str:
        def _clean_descriptor(value: str | None) -> str:
            if value is None:
                return ""

            text = str(value).strip()
            if not text:
                return ""

            text = re.sub(r"^(?:프로젝트|프로그램)\s*개요[:：\-]?\s*", "", text, flags=re.IGNORECASE)
            text = text.replace("\r", " ").replace("\n", " ")
            text = re.sub(r"\s+", " ", text).strip()
            text = re.sub(r"^이\s*(?:프로그램|프로젝트)\s*는\s*", "", text, flags=re.IGNORECASE)
            text = re.sub(r"[.。．]+$", "", text).strip()

            replacements = [
                ("입니다", ""),
                ("이다", ""),
                ("합니다", "하는"),
                ("됩니다", "되는"),
                ("구성됩니다", "구성되는"),
                ("제공합니다", "제공하는"),
                ("제공됩니다", "제공되는"),
                ("포함합니다", "포함하는"),
                ("포함됩니다", "포함되는"),
                ("지원합니다", "지원하는"),
                ("지원됩니다", "지원되는"),
                ("운영됩니다", "운영되는"),
                ("연동됩니다", "연동되는"),
                ("실행됩니다", "실행되는"),
                ("따릅니다", "따르는"),
            ]

            for suffix, replacement in replacements:
                if text.endswith(suffix):
                    text = text[: -len(suffix)] + replacement
                    break

            text = text.strip()
            if text.endswith("다"):
                text = text[:-1].strip()
            if text.endswith("요"):
                text = text[:-1].strip()

            return text.strip()

        def _fallback_descriptor(entries: Sequence[Mapping[str, str]]) -> str:
            for key in ("대분류", "중분류", "소분류"):
                seen: set[str] = set()
                ordered: list[str] = []
                for entry in entries:
                    candidate = str(entry.get(key, "") or "").strip()
                    if not candidate:
                        continue
                    normalized = re.sub(r"\s+", " ", candidate)
                    if normalized and normalized not in seen:
                        seen.add(normalized)
                        ordered.append(normalized)
                if ordered:
                    if len(ordered) == 1:
                        return f"{ordered[0]} 관련"
                    joined = ", ".join(ordered[:3])
                    return f"{joined} 관련"
            return "주요 업무를 지원하는"

        def _compose_sentence(descriptor: str) -> str:
            descriptor = re.sub(r"\s+", " ", descriptor).strip()
            if not descriptor:
                descriptor = "주요 업무를 지원하는"

            suffix_candidates = (
                "프로그램",
                "시스템",
                "플랫폼",
                "솔루션",
                "서비스",
                "애플리케이션",
                "앱",
            )

            if any(descriptor.endswith(suffix) for suffix in suffix_candidates):
                body = descriptor
            else:
                body = f"{descriptor} 프로그램"

            sentence = f"이 프로그램은 {body}이다."
            sentence = re.sub(r"\s+", " ", sentence).strip()
            if not sentence.endswith("."):
                sentence += "."
            return sentence

        def _collect_feature_summaries(entries: Sequence[Mapping[str, str]]) -> list[str]:
            summaries: list[str] = []
            seen: set[str] = set()
            for entry in entries:
                for key in ("기능 설명", "소분류", "중분류", "대분류"):
                    raw_value = entry.get(key, "")
                    if not raw_value:
                        continue
                    candidate = re.sub(r"\s+", " ", str(raw_value).strip())
                    if not candidate:
                        continue
                    if candidate in seen:
                        continue
                    seen.add(candidate)
                    summaries.append(candidate)
                    break
                if len(summaries) >= max_features:
                    break

            if not summaries:
                summaries.append("주요 업무를 지원하는 기능")

            return summaries

        descriptor = _clean_descriptor(raw_overview)
        if not descriptor:
            descriptor = _fallback_descriptor(records)

        first_line = _compose_sentence(descriptor)
        features = _collect_feature_summaries(records)

        lines = [first_line, "기능은"]
        lines.extend(f"- {feature}" for feature in features)
        return "\n".join(lines)

    def _convert_required_documents_to_pdf(
        self,
        uploads: List[BufferedUpload],
        metadata_entries: List[Dict[str, Any]],
    ) -> List[BufferedUpload]:
        if not metadata_entries:
            return uploads

        converted = list(uploads)
        for index, upload in enumerate(uploads):
            metadata = metadata_entries[index] if index < len(metadata_entries) else None
            if not isinstance(metadata, dict):
                continue

            role = str(metadata.get("role") or "").strip().lower()
            if role != "required":
                continue

            doc_id = str(metadata.get("id") or "").strip()
            if doc_id not in {"user-manual", "vendor-feature-list"}:
                continue

            converted_upload = self._convert_single_required_document_to_pdf(
                upload, metadata
            )
            if converted_upload is not upload:
                converted[index] = converted_upload

        return converted

    def _convert_single_required_document_to_pdf(
        self, upload: BufferedUpload, metadata: Dict[str, Any]
    ) -> BufferedUpload:
        extension = self._detect_raw_extension(upload)
        if extension == "pdf":
            return upload

        label = str(
            metadata.get("label")
            or metadata.get("description")
            or metadata.get("id")
            or "문서"
        ).strip()

        original_extension = extension.upper() if extension else ""

        if extension == "docx":
            converted = self._convert_docx_upload_to_pdf(upload, label)
        elif extension == "xlsx":
            converted = self._convert_xlsx_upload_to_pdf(upload, label)
        elif extension == "csv":
            converted = self._convert_csv_upload_to_pdf(upload, label)
        else:
            raise HTTPException(
                status_code=422,
                detail=f"{label}은(는) PDF 또는 지원되는 문서 형식(DOCX, XLSX, CSV)이어야 합니다.",
            )

        self._append_conversion_note(metadata, label, original_extension)
        return converted

    @staticmethod
    def _append_conversion_note(
        metadata: Dict[str, Any], label: str, original_extension: str
    ) -> None:
        if not original_extension:
            return

        note = f"{label} {original_extension} 파일을 PDF로 변환했습니다."
        existing = str(metadata.get("notes") or "").strip()
        metadata["notes"] = f"{existing}\n{note}".strip() if existing else note
        metadata["originalExtension"] = original_extension
        metadata["convertedToPdf"] = True

    @staticmethod
    def _detect_raw_extension(upload: BufferedUpload) -> str:
        extension = os.path.splitext(upload.name)[1].lstrip(".").lower()
        if extension:
            return extension

        content_type = (upload.content_type or "").split(";")[0].strip()
        if content_type:
            guessed = mimetypes.guess_extension(content_type)
            if guessed:
                return guessed.lstrip(".").lower()

        return ""

    @staticmethod
    def _pdf_file_name(original_name: str) -> str:
        base = os.path.splitext(original_name)[0] or "converted"
        return f"{base}.pdf"

    @staticmethod
    def _build_pdf_upload(
        original: BufferedUpload, rows: List[List[str]]
    ) -> BufferedUpload:
        pdf_bytes = AIGenerationService._rows_to_pdf(rows)
        return BufferedUpload(
            name=AIGenerationService._pdf_file_name(original.name),
            content=pdf_bytes,
            content_type="application/pdf",
        )

    @staticmethod
    def _convert_docx_upload_to_pdf(
        upload: BufferedUpload, label: str
    ) -> BufferedUpload:
        try:
            document = Document(io.BytesIO(upload.content))
        except Exception as exc:
            logger.warning(
                "Failed to parse DOCX for PDF conversion; falling back to plain text",
                extra={"label": label, "error": str(exc)},
            )
            decoded = upload.content.decode("utf-8", errors="ignore")
            if not decoded.strip():
                decoded = upload.content.decode("latin-1", errors="ignore")
            lines = decoded.splitlines() or [decoded]
            rows = [[line.strip()] for line in lines]
            return AIGenerationService._build_pdf_upload(upload, rows)

        rows: List[List[str]] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            rows.append([text] if text else [])

        if document.paragraphs and document.tables:
            rows.append([])

        for table in document.tables:
            for row in table.rows:
                cell_values = [cell.text.strip() for cell in row.cells]
                if any(cell_values):
                    rows.append(cell_values)
            rows.append([])

        if rows and not rows[-1]:
            rows.pop()

        return AIGenerationService._build_pdf_upload(upload, rows)

    @staticmethod
    def _convert_xlsx_upload_to_pdf(
        upload: BufferedUpload, label: str
    ) -> BufferedUpload:
        try:
            rows = AIGenerationService._parse_xlsx_rows(upload.content)
        except ValueError as exc:
            raise HTTPException(
                status_code=422,
                detail=f"{label} XLSX 파일을 PDF로 변환하는 중 오류가 발생했습니다.",
            ) from exc

        return AIGenerationService._build_pdf_upload(upload, rows)

    @staticmethod
    def _convert_csv_upload_to_pdf(
        upload: BufferedUpload, label: str
    ) -> BufferedUpload:
        decoded: str | None = None
        last_error: Exception | None = None
        for encoding in ("utf-8-sig", "utf-8", "cp949", "latin-1"):
            try:
                decoded = upload.content.decode(encoding)
                break
            except UnicodeDecodeError as exc:
                last_error = exc

        if decoded is None:
            raise HTTPException(
                status_code=422,
                detail=f"{label} CSV 파일을 읽는 중 오류가 발생했습니다.",
            ) from last_error

        try:
            reader = csv.reader(io.StringIO(decoded))
            rows = [[cell.strip() for cell in row] for row in reader]
        except csv.Error as exc:
            raise HTTPException(
                status_code=422,
                detail=f"{label} CSV 파일을 PDF로 변환하는 중 오류가 발생했습니다.",
            ) from exc

        return AIGenerationService._build_pdf_upload(upload, rows)

    async def generate_csv(
        self,
        project_id: str,
        menu_id: str,
        uploads: List[UploadFile],
        metadata: List[Dict[str, Any]] | None = None,
    ) -> GeneratedCsv:
        try:
            prompt_config = self._prompt_config_service.get_runtime_prompt(menu_id)
        except KeyError as exc:
            raise HTTPException(
                status_code=404, detail="지원하지 않는 생성 메뉴입니다."
            ) from exc

        if not uploads:
            raise HTTPException(status_code=422, detail="업로드된 자료가 없습니다. 파일을 추가해 주세요.")

        buffered: List[BufferedUpload] = []
        for upload in uploads:
            try:
                data = await upload.read()
                name = upload.filename or "업로드된_파일"
                buffered.append(
                    BufferedUpload(
                        name=name,
                        content=data,
                        content_type=upload.content_type,
                    )
                )
            finally:
                await upload.close()

        metadata_entries: List[Dict[str, Any]] = []
        if metadata:
            metadata_entries = [
                dict(entry) if isinstance(entry, dict) else {}
                for entry in metadata
            ]

        buffered = self._convert_required_documents_to_pdf(
            buffered, metadata_entries
        )

        defect_prompt_section: str | None = None
        defect_summary_entries: List[DefectSummaryEntry] | None = None
        defect_image_map: Dict[int, List[BufferedUpload]] = {}

        contexts: List[UploadContext] = []
        for index, upload in enumerate(buffered):
            entry = metadata_entries[index] if index < len(metadata_entries) else None
            contexts.append(UploadContext(upload=upload, metadata=entry))

        contexts.extend(
            self._builtin_attachment_contexts(menu_id, prompt_config.builtin_contexts)
        )

        if menu_id == "defect-report":
            (
                contexts,
                defect_prompt_section,
                defect_summary_entries,
                defect_image_map,
            ) = self._prepare_defect_report_contexts(contexts, prompt_config)

        client = self._get_client()
        uploaded_file_records: List[tuple[str, bool]] = []
        uploaded_attachments: List[AttachmentMetadata] = []

        try:
            context_previews = self._build_context_previews(contexts)
            context_summary = self._context_summary(menu_id, context_previews)

            descriptor_template = (
                prompt_config.attachment_descriptor_template or "{{index}}. {{descriptor}}"
            )
            descriptor_lines: List[str] = []
            for index, preview in enumerate(context_previews, start=1):
                if not preview.include_in_attachment_list:
                    continue
                replacements: Dict[str, str] = {
                    "index": str(index),
                    "descriptor": preview.descriptor,
                    "doc_id": preview.doc_id or "",
                }
                for key, value in preview.metadata.items():
                    replacements[key] = str(value) if value is not None else ""
                line = descriptor_template
                for key, value in replacements.items():
                    line = line.replace(f"{{{{{key}}}}}", value)
                descriptor_lines.append(line.strip())
            descriptor_section = "\n".join(
                line for line in descriptor_lines if line.strip()
            )

            for context in contexts:
                kind = self._attachment_kind(context.upload)
                if kind == "image":
                    image_url = self._image_data_url(context.upload)
                    uploaded_attachments.append(
                        {
                            "kind": "image",
                            "image_url": image_url,
                        }
                    )
                    continue

                file_id = await self._upload_openai_file(client, context)
                metadata_entry = context.metadata or {}
                skip_cleanup = bool(metadata_entry.get("skip_cleanup"))
                uploaded_file_records.append((file_id, skip_cleanup))
                uploaded_attachments.append(
                    {"file_id": file_id, "kind": kind}
                )

            user_prompt_parts: List[str] = []

            base_instruction = prompt_config.user_prompt.strip()
            if base_instruction:
                user_prompt_parts.append(base_instruction)

            for section in prompt_config.user_prompt_sections:
                if not section.enabled:
                    continue
                label = section.label.strip()
                content = section.content.strip()
                if label and content:
                    user_prompt_parts.append(f"{label}\n{content}")
                elif label or content:
                    user_prompt_parts.append(label or content)

            if defect_prompt_section:
                user_prompt_parts.append(defect_prompt_section)

            if contexts:
                heading = prompt_config.scaffolding.attachments_heading.strip()
                intro = prompt_config.scaffolding.attachments_intro.strip()
                if heading:
                    user_prompt_parts.append(heading)
                if intro:
                    user_prompt_parts.append(intro)
                if descriptor_section:
                    user_prompt_parts.append(descriptor_section)

            closing_template = prompt_config.scaffolding.closing_note.strip()
            if closing_template:
                closing_note = closing_template.replace(
                    "{{context_summary}}", context_summary
                ).strip()
            else:
                closing_note = context_summary.strip()
            if closing_note:
                user_prompt_parts.append(closing_note)

            format_warning = prompt_config.scaffolding.format_warning.strip()
            if format_warning:
                user_prompt_parts.append(format_warning)

            user_prompt = "\n\n".join(part for part in user_prompt_parts if part.strip())

            messages = [
                OpenAIMessageBuilder.text_message(
                    "system", prompt_config.system_prompt
                ),
                OpenAIMessageBuilder.text_message(
                    "user",
                    user_prompt,
                    attachments=uploaded_attachments,
                ),
            ]

            normalized_messages = OpenAIMessageBuilder.normalize_messages(messages)

            logger.info(
                "AI generation prompt assembled",
                extra={
                    "project_id": project_id,
                    "menu_id": menu_id,
                    "system_prompt": prompt_config.system_prompt,
                    "user_prompt": user_prompt,
                },
            )

            params = prompt_config.model_parameters
            try:
                response_kwargs: dict[str, object] = {
                    "model": self._settings.openai_model,
                    "input": normalized_messages,
                }

                if params.temperature is not None:
                    response_kwargs["temperature"] = params.temperature
                if params.top_p is not None:
                    response_kwargs["top_p"] = params.top_p
                if params.max_output_tokens is not None:
                    response_kwargs["max_output_tokens"] = (
                        params.max_output_tokens
                    )

                # The Responses API currently rejects presence/frequency penalties.
                # Until OpenAI adds support we simply omit them from the request to
                # avoid TypeError crashes while still honouring other tunables.
                if params.presence_penalty not in (None, 0):
                    logger.warning(
                        "Presence penalty is not supported by the Responses API; "
                        "value will be ignored.",
                        extra={
                            "project_id": project_id,
                            "menu_id": menu_id,
                            "presence_penalty": params.presence_penalty,
                        },
                    )
                if params.frequency_penalty not in (None, 0):
                    logger.warning(
                        "Frequency penalty is not supported by the Responses API; "
                        "value will be ignored.",
                        extra={
                            "project_id": project_id,
                            "menu_id": menu_id,
                            "frequency_penalty": params.frequency_penalty,
                        },
                    )

                response = await asyncio.to_thread(
                    client.responses.create,
                    **response_kwargs,
                )
            except RateLimitError as exc:
                detail = self._format_openai_error(exc)
                raise HTTPException(
                    status_code=429,
                    detail=(
                        "OpenAI 사용량 한도를 초과했습니다. "
                        "관리자에게 문의하거나 잠시 후 다시 시도해 주세요."
                        f" ({detail})"
                    ),
                ) from exc
            except (PermissionDeniedError, BadRequestError, APIError, OpenAIError) as exc:
                detail = self._format_openai_error(exc)
                raise HTTPException(
                    status_code=502,
                    detail=f"OpenAI 호출 중 오류가 발생했습니다: {detail}",
                ) from exc
            except Exception as exc:  # pragma: no cover - 안전망
                logger.exception(
                    "Unexpected error while requesting OpenAI response",
                    extra={"project_id": project_id, "menu_id": menu_id},
                )
                message = str(exc).strip()
                if message:
                    detail = (
                        "OpenAI 응답을 가져오는 중 예기치 않은 오류가 발생했습니다: "
                        f"{message}"
                    )
                else:
                    detail = "OpenAI 응답을 가져오는 중 예기치 않은 오류가 발생했습니다."
                raise HTTPException(status_code=502, detail=detail) from exc

            response_text = self._extract_response_text(response) or ""

            if self._request_log_service is not None:
                try:
                    self._request_log_service.record_request(
                        project_id=project_id,
                        menu_id=menu_id,
                        system_prompt=prompt_config.system_prompt,
                        user_prompt=user_prompt,
                        context_summary=context_summary,
                        response_text=response_text,
                    )
                except Exception:  # pragma: no cover - logging must not fail request
                    logger.exception(
                        "Failed to record prompt request log",
                        extra={"project_id": project_id, "menu_id": menu_id},
                    )

            if not response_text:
                raise HTTPException(status_code=502, detail="OpenAI 응답에서 CSV를 찾을 수 없습니다.")

            sanitized = self._sanitize_csv(response_text)
            project_overview: str | None = None
            if menu_id == "feature-list" and sanitized:
                sanitized, raw_project_overview = self._extract_feature_list_project_overview(
                    sanitized
                )
                records = normalize_feature_list_records(sanitized)
                project_overview = self._format_feature_list_program_overview(
                    records,
                    raw_project_overview,
                )

            if not sanitized:
                raise HTTPException(status_code=502, detail="생성된 CSV 내용이 비어 있습니다.")

            encoded = sanitized.encode("utf-8-sig")
            timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
            safe_project = re.sub(r"[^A-Za-z0-9_-]+", "_", project_id)
            filename = f"{safe_project}_{menu_id}_{timestamp}.csv"

            return GeneratedCsv(
                filename=filename,
                content=encoded,
                csv_text=sanitized,
                defect_summary=defect_summary_entries,
                defect_images=dict(defect_image_map) if defect_image_map else None,
                project_overview=project_overview,
            )
        finally:
            if uploaded_file_records:
                await self._cleanup_openai_files(client, uploaded_file_records)

    @staticmethod
    def _format_openai_error(exc: OpenAIError) -> str:
        message = str(exc).strip()
        details: List[str] = []
        if message:
            details.append(message)

        body = getattr(exc, "body", None)
        if isinstance(body, dict):
            error = body.get("error")
            candidates: List[str] = []
            if isinstance(error, dict):
                for key in ("message", "code", "type"):
                    value = error.get(key)
                    if isinstance(value, str) and value.strip():
                        candidates.append(value.strip())
            elif isinstance(error, str) and error.strip():
                candidates.append(error.strip())

            for value in candidates:
                if value not in details:
                    details.append(value)

        if not details:
            details.append(exc.__class__.__name__)

        return "; ".join(details)

    @staticmethod
    def _extract_response_text(response: Any) -> str | None:
        """Best-effort extraction of the text payload from the Responses API."""

        def _is_non_empty_text(value: object) -> bool:
            return isinstance(value, str) and bool(value.strip())

        text_candidate = getattr(response, "output_text", None)
        if _is_non_empty_text(text_candidate):
            return str(text_candidate)

        containers: List[object] = []
        for attr in ("output", "outputs", "data", "messages"):
            candidate = getattr(response, attr, None)
            if candidate:
                containers.append(candidate)

        if isinstance(response, dict):
            for key in ("output", "outputs", "data", "messages"):
                candidate = response.get(key)
                if candidate:
                    containers.append(candidate)

        for container in containers:
            if isinstance(container, (list, tuple)):
                text = AIGenerationService._extract_from_content(container)
                if text:
                    return text
            elif isinstance(container, dict):
                content = container.get("content")
                if content:
                    normalized = content if isinstance(content, (list, tuple)) else [content]
                    text = AIGenerationService._extract_from_content(normalized)
                    if text:
                        return text
            else:
                content = getattr(container, "content", None)
                if content:
                    normalized = content if isinstance(content, (list, tuple)) else [content]
                    text = AIGenerationService._extract_from_content(normalized)
                    if text:
                        return text

        return None

    @staticmethod
    def _extract_from_content(items: Iterable[object]) -> str | None:
        for item in items:
            content = None
            if isinstance(item, dict):
                content = item.get("content")
            else:
                content = getattr(item, "content", None)

            if not content or isinstance(content, (str, bytes)):
                continue

            for part in content:
                part_type = None
                text_value = None
                if isinstance(part, dict):
                    part_type = part.get("type")
                    text_value = part.get("text")
                else:
                    part_type = getattr(part, "type", None)
                    text_value = getattr(part, "text", None)

                if part_type in {"output_text", "text", "input_text"} and text_value is not None:
                    extracted = text_value
                    if isinstance(text_value, dict):
                        extracted = text_value.get("value")
                    elif hasattr(text_value, "get"):
                        try:
                            extracted = text_value.get("value")  # type: ignore[attr-defined]
                        except Exception:  # pragma: no cover - defensive
                            extracted = text_value

                    text_str = str(extracted).strip() if extracted is not None else ""
                    if text_str:
                        return text_str

        return None

    @staticmethod
    def _image_data_url(upload: BufferedUpload) -> str:
        media_type = (upload.content_type or "").split(";")[0].strip()
        if not media_type:
            guessed, _ = mimetypes.guess_type(upload.name)
            if guessed:
                media_type = guessed

        if not media_type:
            media_type = "application/octet-stream"

        encoded = base64.b64encode(upload.content).decode("ascii")
        return f"data:{media_type};base64,{encoded}"

    def _prepare_defect_report_contexts(
        self, contexts: List[UploadContext], prompt_config: PromptConfig
    ) -> tuple[
        List[UploadContext],
        str | None,
        List[DefectSummaryEntry] | None,
        Dict[int, List[BufferedUpload]],
    ]:
        summary_entries: List[DefectSummaryEntry] | None = None
        prompt_resources: DefectPromptResources | None = None
        image_map: Dict[int, List[BufferedUpload]] = defaultdict(list)
        filtered_contexts: List[UploadContext] = []

        for context in contexts:
            metadata = context.metadata or {}
            upload = context.upload
            defect_index_value = metadata.get("defect_index")
            if defect_index_value is not None:
                try:
                    defect_index = int(defect_index_value)
                except (TypeError, ValueError):
                    defect_index = None
                else:
                    if self._attachment_kind(upload) == "image":
                        image_map[defect_index].append(upload)

            is_json_upload = upload.name.lower().endswith(".json") or (
                (upload.content_type or "").split(";")[0].lower() == "application/json"
            )

            if is_json_upload:
                parsed_entries, parsed_resources = self._parse_defect_summary_upload(upload)
                if summary_entries is None and parsed_entries:
                    summary_entries = parsed_entries
                if prompt_resources is None and parsed_resources:
                    prompt_resources = parsed_resources
                continue

            filtered_contexts.append(context)

        prompt_resources = self._merge_defect_prompt_resources(
            prompt_resources, prompt_config.prompt_resources
        )

        prompt_section: str | None = None
        if (summary_entries and len(summary_entries) > 0) or prompt_resources:
            prompt_section = self._format_defect_prompt_section(
                summary_entries or [], prompt_resources
            )

        return filtered_contexts, prompt_section, summary_entries, image_map

    @staticmethod
    def _merge_defect_prompt_resources(
        parsed: DefectPromptResources | None,
        config_resources: PromptResourcesConfig | None,
    ) -> DefectPromptResources | None:
        if parsed is None and config_resources is None:
            return None

        config_judgement = (
            config_resources.judgement_criteria.strip()
            if config_resources and config_resources.judgement_criteria
            else ""
        )
        config_example = (
            config_resources.output_example.strip()
            if config_resources and config_resources.output_example
            else ""
        )

        parsed_judgement = (
            parsed.judgement_criteria.strip()
            if parsed and parsed.judgement_criteria
            else ""
        )
        parsed_example = (
            parsed.output_example.strip() if parsed and parsed.output_example else ""
        )

        conversation = list(parsed.conversation) if parsed else []

        judgement = parsed_judgement or config_judgement
        example = parsed_example or config_example

        if not judgement and not example and not conversation:
            return None

        return DefectPromptResources(
            judgement_criteria=judgement or None,
            output_example=example or None,
            conversation=conversation,
        )

    @staticmethod
    def _parse_defect_summary_upload(
        upload: BufferedUpload,
    ) -> tuple[List[DefectSummaryEntry], DefectPromptResources | None]:
        try:
            decoded = upload.content.decode("utf-8-sig")
        except UnicodeDecodeError:
            decoded = upload.content.decode("utf-8", errors="ignore")

        try:
            payload = json.loads(decoded)
        except json.JSONDecodeError:
            return [], None

        defects = payload.get("defects") if isinstance(payload, dict) else None
        if not isinstance(defects, list):
            defects = []

        entries: List[DefectSummaryEntry] = []
        for item in defects:
            if not isinstance(item, dict):
                continue
            index_value = item.get("index")
            polished_text = item.get("polishedText")
            if not isinstance(index_value, int) or not isinstance(polished_text, str):
                continue
            original_text = item.get("originalText")
            if not isinstance(original_text, str):
                original_text = ""

            attachments_raw = item.get("attachments")
            attachments: List[DefectSummaryAttachment] = []
            if isinstance(attachments_raw, list):
                for attachment in attachments_raw:
                    if not isinstance(attachment, dict):
                        continue
                    file_name = attachment.get("fileName")
                    original_name = attachment.get("originalFileName")
                    if isinstance(file_name, str) and file_name.strip():
                        attachments.append(
                            DefectSummaryAttachment(
                                file_name=file_name.strip(),
                                original_file_name=original_name
                                if isinstance(original_name, str)
                                else None,
                            )
                        )

            entries.append(
                DefectSummaryEntry(
                    index=index_value,
                    original_text=original_text.strip(),
                    polished_text=polished_text.strip(),
                    attachments=attachments,
                )
            )

        resources_payload = (
            payload.get("promptResources") if isinstance(payload, dict) else None
        )
        prompt_resources: DefectPromptResources | None = None
        if isinstance(resources_payload, dict):
            judgement_raw = resources_payload.get("judgementCriteria")
            judgement = judgement_raw.strip() if isinstance(judgement_raw, str) else None
            example_raw = resources_payload.get("outputExample")
            output_example = example_raw.strip() if isinstance(example_raw, str) else None

            conversation_raw = resources_payload.get("conversation")
            conversation: List[DefectConversationTurn] = []
            if isinstance(conversation_raw, list):
                for entry in conversation_raw:
                    if not isinstance(entry, dict):
                        continue
                    role = entry.get("role")
                    text_raw = entry.get("text")
                    if role not in ("user", "assistant") or not isinstance(text_raw, str):
                        continue
                    text = text_raw.strip()
                    if not text:
                        continue
                    conversation.append(DefectConversationTurn(role=role, text=text))

            if judgement or output_example or conversation:
                prompt_resources = DefectPromptResources(
                    judgement_criteria=judgement,
                    output_example=output_example,
                    conversation=conversation,
                )

        return entries, prompt_resources

    @staticmethod
    def _format_defect_prompt_section(
        entries: List[DefectSummaryEntry],
        resources: DefectPromptResources | None = None,
    ) -> str | None:
        lines: List[str] = []

        if entries:
            lines.append("정제된 결함 목록")
            lines.append("")
            for entry in sorted(entries, key=lambda item: item.index):
                polished = entry.polished_text or "-"
                lines.append(f"{entry.index}. {polished}")
                if entry.original_text:
                    lines.append(f"   - 원문: {entry.original_text}")
                if entry.attachments:
                    names = ", ".join(att.file_name for att in entry.attachments)
                    lines.append(f"   - 첨부 이미지: {names}")

        if resources:
            def add_section(title: str, body: str | List[str]) -> None:
                if isinstance(body, list):
                    content_lines = [line for line in body if line.strip()]
                else:
                    text = body.strip()
                    if not text:
                        return
                    content_lines = [text]
                if not content_lines:
                    return
                if lines and lines[-1] != "":
                    lines.append("")
                lines.append(title)
                lines.append("")
                lines.extend(content_lines)

            if resources.judgement_criteria:
                add_section("결함 판단 기준", resources.judgement_criteria)
            if resources.output_example:
                add_section("출력 예시", resources.output_example)
            if resources.conversation:
                conversation_lines: List[str] = []
                for index, turn in enumerate(resources.conversation, start=1):
                    text = turn.text.strip()
                    if not text:
                        continue
                    speaker = "사용자" if turn.role == "user" else "GPT"
                    conversation_lines.append(f"{index}. {speaker}: {text}")
                add_section("이전 대화", conversation_lines)

        if not lines:
            return None

        return "\n".join(lines).strip()

    def _builtin_attachment_contexts(
        self, menu_id: str, builtin_contexts: List[PromptBuiltinContext]
    ) -> List[UploadContext]:
        contexts: List[UploadContext] = []
        for builtin in builtin_contexts:
            if not builtin.include_in_prompt:
                continue
            upload = self._load_builtin_upload(menu_id, builtin)
            metadata: Dict[str, Any] = {
                "role": "additional",
                "label": builtin.label,
                "description": builtin.description,
                "source_path": builtin.source_path,
                "show_in_attachment_list": builtin.show_in_attachment_list,
                "skip_cleanup": True,
            }
            contexts.append(UploadContext(upload=upload, metadata=metadata))
        return contexts

    def _locate_builtin_source(self, path_hint: str) -> tuple[Path | None, List[Path]]:
        """Resolve the on-disk path for a builtin attachment.

        Historically the API server has been executed from different working
        directories (package install, repo checkout, Docker image).  In those
        environments the template assets may live under either the backend
        package directory (e.g. ``backend/template``) or directly under the
        application root (``/app/template`` inside the container).  The
        original implementation assumed a single location which caused
        ``FileNotFoundError`` when the active runtime layout differed,
        surfacing to the user as “내장 XLSX 템플릿을 찾을 수 없습니다.”.

        To make the lookup resilient we try several sensible base paths and
        keep track of the attempted locations for diagnostics.
        """

        attempted: List[Path] = []
        requested = Path(path_hint)

        def _candidate(path: Path) -> Optional[Path]:
            resolved = path if path.is_absolute() else path.resolve()
            attempted.append(resolved)
            if resolved.exists() and resolved.is_file():
                return resolved
            return None

        if requested.is_absolute():
            resolved = _candidate(requested)
            if resolved is not None:
                return resolved, attempted

        override_root = getattr(self._settings, "builtin_template_root", None)
        if override_root:
            override_path = Path(override_root)

            if override_path.is_file():
                resolved = _candidate(override_path)
                if resolved is not None:
                    return resolved, attempted

            relative_variants: List[Path] = []

            if str(requested):
                relative_variants.append(requested)

            try:
                relative_to_template = requested.relative_to(Path("template"))
            except ValueError:
                relative_to_template = None
            if relative_to_template and str(relative_to_template):
                relative_variants.append(relative_to_template)

            parts = list(requested.parts)
            if parts and parts[0] == "backend":
                relative_variants.append(Path(*parts[1:]))
                if len(parts) > 1 and parts[1] == "template":
                    relative_variants.append(Path(*parts[2:]))

            seen: set[Path] = set()
            ordered_variants: List[Path] = []
            for variant in relative_variants:
                variant_str = str(variant)
                if not variant_str or variant_str == ".":
                    continue
                if variant in seen:
                    continue
                seen.add(variant)
                ordered_variants.append(variant)

            if override_path.is_dir():
                for variant in ordered_variants:
                    resolved = _candidate(override_path / variant)
                    if resolved is not None:
                        return resolved, attempted
            else:
                for variant in ordered_variants:
                    resolved = _candidate(override_path.parent / variant)
                    if resolved is not None:
                        return resolved, attempted

        base_path = Path(__file__).resolve().parents[2]
        resolved = _candidate(base_path / requested)
        if resolved is not None:
            return resolved, attempted

        repo_root = base_path.parent
        if repo_root != base_path:
            resolved = _candidate(repo_root / requested)
            if resolved is not None:
                return resolved, attempted

        template_root = base_path / "template"
        if template_root.exists():
            try:
                relative_to_template = requested.relative_to(Path("template"))
            except ValueError:
                relative_to_template = requested

            resolved = _candidate(template_root / relative_to_template)
            if resolved is not None:
                return resolved, attempted

            # As a last resort search by filename inside the template tree so
            # renamed folders (e.g. when the repo is vendored) still resolve.
            if requested.name:
                for match in template_root.rglob(requested.name):
                    resolved = _candidate(match)
                    if resolved is not None:
                        return resolved, attempted

        return None, attempted

    def _load_builtin_upload(
        self, menu_id: str, builtin: PromptBuiltinContext
    ) -> BufferedUpload:
        source_path, attempted_paths = self._locate_builtin_source(builtin.source_path)
        if source_path is None:
            logger.error(
                "내장 컨텍스트 파일을 찾을 수 없습니다.",
                extra={
                    "menu_id": menu_id,
                    "path": builtin.source_path,
                    "label": builtin.label,
                    "attempted_paths": [str(path) for path in attempted_paths],
                },
            )
            raise HTTPException(
                status_code=500,
                detail="내장 컨텍스트 파일을 찾을 수 없습니다.",
            )
        if builtin.render_mode == "xlsx-to-pdf":
            return self._load_xlsx_as_pdf(menu_id, source_path, builtin.label)

        try:
            content = source_path.read_bytes()
        except FileNotFoundError as exc:
            logger.error(
                "내장 컨텍스트 파일을 찾을 수 없습니다.",
                extra={
                    "menu_id": menu_id,
                    "path": str(source_path),
                    "label": builtin.label,
                },
            )
            raise HTTPException(
                status_code=500,
                detail="내장 컨텍스트 파일을 찾을 수 없습니다.",
            ) from exc
        except OSError as exc:
            logger.error(
                "내장 컨텍스트 파일을 읽는 중 오류가 발생했습니다.",
                extra={
                    "menu_id": menu_id,
                    "path": str(source_path),
                    "label": builtin.label,
                    "error": str(exc),
                },
            )
            raise HTTPException(
                status_code=500,
                detail="내장 컨텍스트 파일을 읽는 중 오류가 발생했습니다.",
            ) from exc

        guessed_type, _ = mimetypes.guess_type(source_path.name)
        if builtin.render_mode == "text":
            content_type = "text/plain; charset=utf-8"
        elif builtin.render_mode == "image":
            content_type = guessed_type or "image/png"
        else:
            content_type = guessed_type or "application/octet-stream"

        return BufferedUpload(
            name=source_path.name,
            content=content,
            content_type=content_type,
        )

    @staticmethod
    def _load_xlsx_as_pdf(menu_id: str, template_path: Path, label: str) -> BufferedUpload:
        try:
            content = template_path.read_bytes()
        except FileNotFoundError as exc:
            logger.error(
                "내장 XLSX 템플릿을 찾을 수 없습니다.",
                extra={
                    "menu_id": menu_id,
                    "path": str(template_path),
                    "label": label,
                },
            )
            raise HTTPException(
                status_code=500,
                detail="내장 XLSX 템플릿을 찾을 수 없습니다.",
            ) from exc
        except OSError as exc:
            logger.error(
                "내장 XLSX 템플릿을 읽는 중 오류가 발생했습니다.",
                extra={
                    "menu_id": menu_id,
                    "path": str(template_path),
                    "label": label,
                    "error": str(exc),
                },
            )
            raise HTTPException(
                status_code=500,
                detail="내장 XLSX 템플릿을 읽는 중 오류가 발생했습니다.",
            ) from exc

        try:
            rows = AIGenerationService._parse_xlsx_rows(content)
        except ValueError as exc:
            logger.error(
                "내장 XLSX 템플릿을 PDF로 변환하는 중 오류가 발생했습니다.",
                extra={
                    "menu_id": menu_id,
                    "path": str(template_path),
                    "label": label,
                    "error": str(exc),
                },
            )
            raise HTTPException(
                status_code=500,
                detail="내장 XLSX 템플릿을 PDF로 변환하는 중 오류가 발생했습니다.",
            ) from exc

        pdf_bytes = AIGenerationService._rows_to_pdf(rows)

        return BufferedUpload(
            name=template_path.with_suffix(".pdf").name,
            content=pdf_bytes,
            content_type="application/pdf",
        )

    @staticmethod
    def _rows_to_pdf(rows: List[List[str]]) -> bytes:
        lines: List[str] = []
        for row in rows:
            if row:
                line = ", ".join(cell.strip() for cell in row)
            else:
                line = ""
            lines.append(line)

        return AIGenerationService._lines_to_pdf(lines)

    @staticmethod
    def _lines_to_pdf(lines: List[str]) -> bytes:
        if not lines:
            lines = [""]

        def _escape(text: str) -> str:
            encoded = ("\ufeff" + text).encode("utf-16-be")
            return "".join(f"\\{byte:03o}" for byte in encoded)

        content_lines = [
            "BT",
            "/F1 11 Tf",
            "1 0 0 1 72 770 Tm",
            "14 TL",
        ]
        for line in lines:
            escaped = _escape(line)
            content_lines.append(f"({escaped}) Tj")
            content_lines.append("T*")
        content_lines.append("ET")

        content_stream = "\n".join(content_lines).encode("utf-8")

        objects: List[bytes] = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 6 0 R /Resources << /Font << /F1 4 0 R >> >> >>",
            b"<< /Type /Font /Subtype /Type0 /BaseFont /HYGoThic-Medium /Encoding /UniKS-UCS2-H /DescendantFonts [5 0 R] >>",
            b"<< /Type /Font /Subtype /CIDFontType0 /BaseFont /HYGoThic-Medium /CIDSystemInfo << /Registry (Adobe) /Ordering (Korea1) /Supplement 0 >> /DW 1000 >>",
            b"<< /Length %d >>\nstream\n" % len(content_stream)
            + content_stream
            + b"\nendstream",
        ]

        buffer = io.BytesIO()
        buffer.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets: List[int] = []
        for index, obj in enumerate(objects, start=1):
            offsets.append(buffer.tell())
            buffer.write(f"{index} 0 obj\n".encode("ascii"))
            buffer.write(obj)
            buffer.write(b"\nendobj\n")

        xref_offset = buffer.tell()
        total_objects = len(objects) + 1
        buffer.write(f"xref\n0 {total_objects}\n".encode("ascii"))
        buffer.write(b"0000000000 65535 f \n")
        for offset in offsets:
            buffer.write(f"{offset:010d} 00000 n \n".encode("ascii"))
        buffer.write(
            b"trailer\n<< /Size "
            + str(total_objects).encode("ascii")
            + b" /Root 1 0 R >>\nstartxref\n"
            + str(xref_offset).encode("ascii")
            + b"\n%%EOF\n"
        )

        return buffer.getvalue()

    @staticmethod
    def _parse_xlsx_rows(content: bytes) -> List[List[str]]:
        try:
            archive = zipfile.ZipFile(io.BytesIO(content))
        except zipfile.BadZipFile as exc:
            raise ValueError("잘못된 XLSX 형식입니다.") from exc

        with archive:
            shared_strings = AIGenerationService._read_shared_strings(archive)
            try:
                with archive.open("xl/worksheets/sheet1.xml") as sheet_file:
                    tree = ET.parse(sheet_file)
            except KeyError as exc:
                raise ValueError("기본 시트를 찾을 수 없습니다.") from exc
            except ET.ParseError as exc:
                raise ValueError("시트 XML을 해석할 수 없습니다.") from exc

            namespace = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
            root = tree.getroot()
            sheet_data = root.find("main:sheetData", namespace)
            if sheet_data is None:
                return []

            rows: List[List[str]] = []
            for row_elem in sheet_data.findall("main:row", namespace):
                row_values: List[str] = []
                for cell_elem in row_elem.findall("main:c", namespace):
                    column_index = AIGenerationService._column_index_from_ref(cell_elem.get("r"))
                    value = AIGenerationService._extract_cell_value(cell_elem, shared_strings, namespace)
                    if column_index is None:
                        column_index = len(row_values)
                    while len(row_values) <= column_index:
                        row_values.append("")
                    row_values[column_index] = value
                rows.append(row_values)

            return rows

    @staticmethod
    def _read_shared_strings(archive: zipfile.ZipFile) -> List[str]:
        try:
            with archive.open("xl/sharedStrings.xml") as handle:
                tree = ET.parse(handle)
        except KeyError:
            return []
        except ET.ParseError as exc:
            raise ValueError("공유 문자열 XML을 해석할 수 없습니다.") from exc

        namespace = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        strings: List[str] = []
        root = tree.getroot()
        for si in root.findall("main:si", namespace):
            text_parts = [node.text or "" for node in si.findall(".//main:t", namespace)]
            strings.append("".join(text_parts))
        return strings

    @staticmethod
    def _column_index_from_ref(ref: str | None) -> int | None:
        if not ref:
            return None
        match = re.match(r"([A-Z]+)", ref)
        if not match:
            return None
        letters = match.group(1)
        index = 0
        for letter in letters:
            index = index * 26 + (ord(letter) - ord("A") + 1)
        return index - 1

    @staticmethod
    def _extract_cell_value(
        cell_elem: ET.Element,
        shared_strings: List[str],
        namespace: Dict[str, str],
    ) -> str:
        cell_type = cell_elem.get("t")
        if cell_type == "s":
            index_text = cell_elem.findtext("main:v", default="", namespaces=namespace)
            try:
                shared_index = int(index_text)
            except (TypeError, ValueError):
                return ""
            if 0 <= shared_index < len(shared_strings):
                return shared_strings[shared_index]
            return ""

        if cell_type == "inlineStr":
            text_nodes = cell_elem.findall(".//main:t", namespace)
            return "".join(node.text or "" for node in text_nodes)

        value = cell_elem.findtext("main:v", default="", namespaces=namespace)
        if value:
            return value

        text_nodes = cell_elem.findall(".//main:t", namespace)
        if text_nodes:
            return "".join(node.text or "" for node in text_nodes)

        return ""
