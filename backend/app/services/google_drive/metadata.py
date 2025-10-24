"""Utilities for extracting project metadata from DOCX documents."""
from __future__ import annotations

import io
import re
from typing import Dict, Iterable, Optional

from docx import Document
from fastapi import HTTPException

EXAM_NUMBER_PATTERN = re.compile(r"GS-[A-Z]-\d{2}-\d{4}")

__all__ = [
    "normalize_label",
    "extract_project_metadata",
    "build_project_folder_name",
]


def normalize_label(value: str) -> str:
    return re.sub(r"\s+", "", value or "")


def extract_project_metadata(file_bytes: bytes) -> Dict[str, str]:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:  # pragma: no cover - library level validation
        raise HTTPException(status_code=422, detail="시험 합의서 파일을 읽지 못했습니다.") from exc

    exam_number: Optional[str] = None
    company_name: Optional[str] = None
    product_name: Optional[str] = None

    def _extract_from_cells(cells: Iterable[str]) -> None:
        nonlocal exam_number, company_name, product_name
        cell_iter = iter(cells)
        for label, value in zip(cell_iter, cell_iter):
            normalized_label = normalize_label(label)
            stripped_value = value.strip()
            if not stripped_value:
                continue
            if normalized_label == "시험신청번호":
                match = EXAM_NUMBER_PATTERN.search(stripped_value)
                if match:
                    exam_number = match.group(0)
            elif normalized_label == "제조자":
                company_name = stripped_value
            elif normalized_label.startswith("제품명및버전"):
                lines = [line.strip() for line in stripped_value.split("\n") if line.strip()]
                if lines:
                    last_line = lines[-1]
                    if ":" in last_line:
                        product_name = last_line.split(":", 1)[1].strip()
                    else:
                        product_name = last_line

    for table in document.tables:
        cells: list[str] = []
        for row in table.rows:
            if len(row.cells) >= 2:
                if row.cells[0].text.strip() and row.cells[1].text.strip():
                    cells.append(row.cells[0].text.strip())
                    cells.append(row.cells[1].text.strip())
                if len(row.cells) >= 4 and row.cells[2].text.strip() and row.cells[3].text.strip():
                    cells.append(row.cells[2].text.strip())
                    cells.append(row.cells[3].text.strip())

        if cells:
            _extract_from_cells(cells)

    if exam_number is None:
        combined_text = "\n".join(
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text and paragraph.text.strip()
        )
        match = EXAM_NUMBER_PATTERN.search(combined_text)
        if match:
            exam_number = match.group(0)

    if not exam_number:
        raise HTTPException(status_code=422, detail="시험신청 번호를 찾을 수 없습니다.")

    if not company_name:
        raise HTTPException(status_code=422, detail="제조자(업체명)를 찾을 수 없습니다.")

    if not product_name:
        raise HTTPException(status_code=422, detail="제품명 및 버전을 찾을 수 없습니다.")

    return {
        "exam_number": exam_number.strip(),
        "company_name": company_name.strip(),
        "product_name": product_name.strip(),
    }


def build_project_folder_name(metadata: Dict[str, str]) -> str:
    exam_number = metadata.get("exam_number", "").strip()
    company_name = metadata.get("company_name", "").strip()
    product_name = metadata.get("product_name", "").strip()
    return f"[{exam_number}] {company_name} - {product_name}"
