from __future__ import annotations

import csv
import io
import json
import logging
import mimetypes
import os
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple, TypedDict

import httpx
from docx import Document
from fastapi import HTTPException, UploadFile
from openpyxl import Workbook, load_workbook

from ..config import Settings
from ..token_store import StoredTokens, TokenStorage
from . import excel_templates as excel_templates_service
from .excel_templates import (
    FEATURE_LIST_EXPECTED_HEADERS,
    match_feature_list_header,
    populate_defect_report,
    populate_feature_list,
    populate_security_report,
    populate_testcase_list,
)

if "extract_feature_list_overview" not in globals():
    extract_feature_list_overview = excel_templates_service.extract_feature_list_overview
from .oauth import GOOGLE_TOKEN_ENDPOINT, GoogleOAuthService

logger = logging.getLogger(__name__)

DRIVE_API_BASE = "https://www.googleapis.com/drive/v3"
DRIVE_FILES_ENDPOINT = "/files"
DRIVE_UPLOAD_BASE = "https://www.googleapis.com/upload/drive/v3"
DRIVE_FOLDER_MIME_TYPE = "application/vnd.google-apps.folder"
GOOGLE_SHEETS_MIME_TYPE = "application/vnd.google-apps.spreadsheet"
XLSX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
PROJECT_FOLDER_BASE_NAME = "GS-X-X-XXXX"
TEMPLATE_ROOT = Path(__file__).resolve().parents[2] / "template"
PLACEHOLDER_PATTERNS: Tuple[str, ...] = (
    "GS-B-XX-XXXX",
    "GS-B-2X-XXXX",
    "GS-X-X-XXXX",
)
EXAM_NUMBER_PATTERN = re.compile(r"GS-[A-Z]-\d{2}-\d{4}")


_SHARED_CRITERIA_FILE_CANDIDATES: Tuple[str, ...] = (
    "보안성 결함판단기준표 v1.0.xlsx",
    "결함판단기준표 v1.0.xlsx",
    "결함 판단 기준표 v1.0.xlsx",
    "결함 판단기준표 v1.0.xlsx",
    "공유 결함판단기준표 v1.0.xlsx",
    "공유 결함 판단 기준표 v1.0.xlsx",
)


def _normalize_shared_criteria_name(name: str) -> str:
    base = name.strip().lower()
    if base.endswith(".xlsx"):
        base = base[:-5]
    return re.sub(r"\s+", "", base)


_SHARED_CRITERIA_NORMALIZED_NAMES = {
    _normalize_shared_criteria_name(candidate) for candidate in _SHARED_CRITERIA_FILE_CANDIDATES
}
_PREFERRED_SHARED_CRITERIA_FILE_NAME = _SHARED_CRITERIA_FILE_CANDIDATES[0]


def _normalize_drive_text(value: str) -> str:
    normalized = unicodedata.normalize("NFKC", value or "")
    normalized = normalized.replace("\xa0", " ")
    normalized = normalized.strip().lower()
    return re.sub(r"\s+", " ", normalized)


def _squash_drive_text(value: str) -> str:
    if not value:
        return ""
    return re.sub(r"[\s._\-()]+", "", value)


def _strip_drive_extension(value: str) -> str:
    if "." in value:
        return value.rsplit(".", 1)[0]
    return value


def _strip_drive_version_suffix(value: str) -> str:
    return re.sub(r"v\s*\d+(?:[._\-]\d+)*$", "", value).strip()


def _drive_name_variants(value: str) -> Tuple[str, ...]:
    normalized = _normalize_drive_text(value)
    if not normalized:
        return tuple()

    variants = {normalized}

    squashed = _squash_drive_text(normalized)
    if squashed:
        variants.add(squashed)

    stem = _strip_drive_extension(normalized)
    if stem and stem != normalized:
        variants.add(stem)
        squashed_stem = _squash_drive_text(stem)
        if squashed_stem:
            variants.add(squashed_stem)

    versionless = _strip_drive_version_suffix(stem)
    if versionless and versionless not in variants:
        variants.add(versionless)
        squashed_versionless = _squash_drive_text(versionless)
        if squashed_versionless:
            variants.add(squashed_versionless)

    return tuple(variant for variant in variants if len(variant) >= 2)


def _drive_name_matches(value: str, expected: str) -> bool:
    actual_tokens = set(_drive_name_variants(value))
    expected_tokens = set(_drive_name_variants(expected))
    if not actual_tokens or not expected_tokens:
        return False
    return bool(actual_tokens & expected_tokens)


def _drive_suffix_matches(name: str, suffix: str) -> bool:
    if not suffix:
        return False
    suffix_tokens = set(_drive_name_variants(suffix))
    if not suffix_tokens:
        return False

    name_tokens = set(_drive_name_variants(name))
    if not name_tokens:
        return False

    for token in name_tokens:
        for suffix_token in suffix_tokens:
            if suffix_token and (token.endswith(suffix_token) or suffix_token in token):
                return True
    return False


def _looks_like_header_row(values: Sequence[Any], expected: Sequence[str]) -> bool:
    if not values:
        return False

    normalized_values = [
        _normalize_drive_text(str(value)) if value is not None else ""
        for value in values
    ]
    squashed_values = [_squash_drive_text(value) for value in normalized_values]
    normalized_expected = [_normalize_drive_text(name) for name in expected]
    squashed_expected = [_squash_drive_text(name) for name in normalized_expected]

    matches = 0
    for expected_value, expected_squashed in zip(normalized_expected, squashed_expected):
        if not expected_value and not expected_squashed:
            continue

        for actual_value, actual_squashed in zip(normalized_values, squashed_values):
            if not actual_value and not actual_squashed:
                continue

            normalized_match = (
                bool(expected_value)
                and bool(actual_value)
                and (
                    actual_value == expected_value
                    or expected_value in actual_value
                    or actual_value in expected_value
                )
            )
            squashed_match = (
                bool(expected_squashed)
                and bool(actual_squashed)
                and expected_squashed in actual_squashed
            )

            if normalized_match or squashed_match:
                matches += 1
                break

    if not matches:
        return False

    threshold = max(1, len(normalized_expected) - 1)
    return matches >= threshold


def _is_shared_criteria_candidate(filename: str) -> bool:
    """
    템플릿 파일명이 공유 결함판단기준표 후보들과 동일(공백/대소문자/확장자 무시)한지 판정.
    프로젝트 폴더 복사에서 제외하기 위해 사용.
    """
    try:
        normalized = _normalize_shared_criteria_name(filename)
    except Exception:
        return False
    return normalized in _SHARED_CRITERIA_NORMALIZED_NAMES


class _SpreadsheetRule(TypedDict):
    folder_name: str
    file_suffix: str
    populate: Any


_SPREADSHEET_RULES: Dict[str, _SpreadsheetRule] = {
    "feature-list": {
        "folder_name": "가.계획",
        "file_suffix": "기능리스트 v1.0.xlsx",
        "populate": populate_feature_list,
    },
    "testcase-generation": {
        "folder_name": "나.설계",
        "file_suffix": "테스트케이스.xlsx",
        "populate": populate_testcase_list,
    },
    "defect-report": {
        "folder_name": "다.수행",
        "file_suffix": "결함리포트 v1.0.xlsx",
        "populate": populate_defect_report,
    },
    "security-report": {
        "folder_name": "다.수행",
        "file_suffix": "결함리포트 v1.0.xlsx",
        "populate": populate_security_report,
    },
}


@dataclass
class _ResolvedSpreadsheet:
    rule: _SpreadsheetRule
    tokens: StoredTokens
    folder_id: str
    file_id: str
    file_name: str
    mime_type: Optional[str]
    modified_time: Optional[str]
    content: Optional[bytes] = None


_FEATURE_LIST_START_ROW = 8
_FEATURE_LIST_SHEET_CANDIDATES: Tuple[str, ...] = (
    "기능리스트",
    "기능 리스트",
    "feature list",
)


class GoogleDriveService:
    """High level operations for interacting with Google Drive."""

    def __init__(
        self,
        settings: Settings,
        token_storage: TokenStorage,
        oauth_service: GoogleOAuthService,
    ) -> None:
        self._settings = settings
        self._token_storage = token_storage
        self._oauth_service = oauth_service

    @staticmethod
    def _normalize_label(value: str) -> str:
        return re.sub(r"\s+", "", value or "")

    @staticmethod
    def _extract_project_metadata(file_bytes: bytes) -> Dict[str, str]:
        try:
            document = Document(io.BytesIO(file_bytes))
        except Exception as exc:  # pragma: no cover - library level validation
            raise HTTPException(status_code=422, detail="시험 합의서 파일을 읽지 못했습니다.") from exc

        exam_number: Optional[str] = None
        company_name: Optional[str] = None
        product_name: Optional[str] = None

        def _extract_from_cells(cells: Iterable[str]) -> None:
            nonlocal exam_number, company_name, product_name
            cell_iter = iter(cells)
            for label, value in zip(cell_iter, cell_iter):
                normalized_label = GoogleDriveService._normalize_label(label)
                stripped_value = value.strip()
                if not stripped_value:
                    continue
                if normalized_label == "시험신청번호":
                    match = EXAM_NUMBER_PATTERN.search(stripped_value)
                    if match:
                        exam_number = match.group(0)
                elif normalized_label == "제조자":
                    company_name = stripped_value
                elif normalized_label.startswith("제품명및버전"):
                    lines = [line.strip() for line in stripped_value.split('\n') if line.strip()]
                    if lines:
                        last_line = lines[-1]
                        if ":" in last_line:
                            product_name = last_line.split(":", 1)[1].strip()
                        else:
                            product_name = last_line

        for table in document.tables:
            cells: List[str] = []
            for row in table.rows:
                if len(row.cells) >= 2:
                    if row.cells[0].text.strip() and row.cells[1].text.strip():
                         cells.append(row.cells[0].text.strip())
                         cells.append(row.cells[1].text.strip())
                    if len(row.cells) >= 4 and row.cells[2].text.strip() and row.cells[3].text.strip():
                         cells.append(row.cells[2].text.strip())
                         cells.append(row.cells[3].text.strip())

            if cells:
                _extract_from_cells(cells)

        if exam_number is None:
            combined_text = "\n".join(
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text and paragraph.text.strip()
            )
            match = EXAM_NUMBER_PATTERN.search(combined_text)
            if match:
                exam_number = match.group(0)

        if not exam_number:
            raise HTTPException(status_code=422, detail="시험신청 번호를 찾을 수 없습니다.")

        if not company_name:
            raise HTTPException(status_code=422, detail="제조자(업체명)를 찾을 수 없습니다.")

        if not product_name:
            raise HTTPException(status_code=422, detail="제품명 및 버전을 찾을 수 없습니다.")

        return {
            "exam_number": exam_number.strip(),
            "company_name": company_name.strip(),
            "product_name": product_name.strip(),
        }

    @staticmethod
    def _build_project_folder_name(metadata: Dict[str, str]) -> str:
        exam_number = metadata.get("exam_number", "").strip()
        company_name = metadata.get("company_name", "").strip()
        product_name = metadata.get("product_name", "").strip()

        return f"[{exam_number}] {company_name} - {product_name}"

    @staticmethod
    def _replace_placeholders(text: str, exam_number: str) -> str:
        result = text
        for placeholder in PLACEHOLDER_PATTERNS:
            result = result.replace(placeholder, exam_number)
        return result

    @staticmethod
    def _prepare_template_file_content(path: Path, exam_number: str) -> bytes:
        raw_bytes = path.read_bytes()
        extension = path.suffix.lower()
        if extension in {".docx", ".xlsx", ".pptx"}:
            raw_bytes = GoogleDriveService._replace_in_office_document(raw_bytes, exam_number)
        return raw_bytes

    @staticmethod
    def _replace_in_office_document(data: bytes, exam_number: str) -> bytes:
        original = io.BytesIO(data)
        updated = io.BytesIO()
        with zipfile.ZipFile(original, "r") as source_zip:
            with zipfile.ZipFile(updated, "w") as target_zip:
                for item in source_zip.infolist():
                    content = source_zip.read(item.filename)
                    try:
                        decoded = content.decode("utf-8")
                    except UnicodeDecodeError:
                        target_zip.writestr(item, content)
                        continue
                    replaced = GoogleDriveService._replace_placeholders(decoded, exam_number)
                    target_zip.writestr(item, replaced.encode("utf-8"))
        return updated.getvalue()

    @staticmethod
    def _guess_mime_type(path: Path) -> str:
        mime_type, _ = mimetypes.guess_type(path.name)
        return mime_type or "application/octet-stream"

    @staticmethod
    def _build_default_shared_criteria_workbook() -> bytes:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "결함판단기준"
        headers = [
            "Invicti 결과",
            "결함 요약",
            "결함정도",
            "발생빈도",
            "품질특성",
            "결함 설명",
            "결함 제외 여부",
        ]
        sheet.append(headers)
        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _load_shared_criteria_template_bytes() -> bytes:
        for candidate in _SHARED_CRITERIA_FILE_CANDIDATES:
            template_path = TEMPLATE_ROOT / candidate
            if template_path.exists():
                return template_path.read_bytes()
        logger.warning(
            "Shared criteria template not found in template folder; generating a default workbook."
        )
        return GoogleDriveService._build_default_shared_criteria_workbook()

    async def _copy_template_to_drive(
        self,
        tokens: StoredTokens,
        *,
        parent_id: str,
        exam_number: str,
    ) -> StoredTokens:
        if not TEMPLATE_ROOT.exists():
            raise HTTPException(status_code=500, detail="template 폴더를 찾을 수 없습니다.")

        path_to_folder_id: Dict[Path, str] = {TEMPLATE_ROOT: parent_id}
        for root_dir, dirnames, filenames in os.walk(TEMPLATE_ROOT):
            current_path = Path(root_dir)
            drive_parent_id = path_to_folder_id[current_path]

            for dirname in sorted(dirnames):
                local_dir = current_path / dirname
                folder_name = self._replace_placeholders(dirname, exam_number)
                folder, tokens = await self._create_child_folder(
                    tokens,
                    name=folder_name,
                    parent_id=drive_parent_id,
                )
                path_to_folder_id[local_dir] = str(folder["id"])

            for filename in sorted(filenames):
                # ✅ 공유 결함판단기준표는 프로젝트 폴더에 복사하지 않음
                if _is_shared_criteria_candidate(filename):
                    logger.info("Skip copying shared criteria into project: %s", filename)
                    continue

                local_file = current_path / filename
                target_name = self._replace_placeholders(filename, exam_number)
                content = self._prepare_template_file_content(local_file, exam_number)
                mime_type = self._guess_mime_type(local_file)
                _, tokens = await self._upload_file_to_folder(
                    tokens,
                    file_name=target_name,
                    parent_id=drive_parent_id,
                    content=content,
                    content_type=mime_type,
                )

        return tokens

    def _load_tokens(self, google_id: Optional[str]) -> StoredTokens:
        if google_id:
            stored = self._token_storage.load_by_google_id(google_id)
            if stored is None:
                raise HTTPException(status_code=404, detail="요청한 Google 계정 토큰을 찾을 수 없습니다.")
            return stored

        accounts = self._token_storage.list_accounts()
        if not accounts:
            raise HTTPException(status_code=404, detail="저장된 Google 계정이 없습니다. 먼저 로그인하세요.")

        for account in accounts:
            stored = self._token_storage.load_by_google_id(account.google_id)
            if stored is not None:
                return stored

        raise HTTPException(status_code=404, detail="저장된 Google 계정 토큰을 찾을 수 없습니다.")

    def _is_token_expired(self, tokens: StoredTokens) -> bool:
        if tokens.expires_in <= 0:
            return False

        expires_at = tokens.saved_at + timedelta(seconds=tokens.expires_in)
        now = datetime.now(timezone.utc)
        return now >= expires_at - timedelta(minutes=1)

    async def _refresh_access_token(self, tokens: StoredTokens) -> StoredTokens:
        if not tokens.refresh_token:
            raise HTTPException(status_code=401, detail="Google 인증이 만료되었습니다. 다시 로그인해주세요.")

        data = {
            "client_id": self._settings.client_id,
            "client_secret": self._settings.client_secret,
            "refresh_token": tokens.refresh_token,
            "grant_type": "refresh_token",
        }

        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.post(GOOGLE_TOKEN_ENDPOINT, data=data)

        if response.is_error:
            logger.error("Google token refresh failed: %s", response.text)
            raise HTTPException(status_code=502, detail="Google 토큰을 새로고침하지 못했습니다. 다시 로그인해주세요.")

        payload = response.json()
        access_token = payload.get("access_token")
        if not isinstance(access_token, str) or not access_token:
            logger.error("Google token refresh response missing access_token: %s", payload)
            raise HTTPException(status_code=502, detail="Google 토큰을 새로고침하지 못했습니다. 다시 로그인해주세요.")

        merged_payload: Dict[str, Any] = {
            "access_token": access_token,
            "refresh_token": payload.get("refresh_token") or tokens.refresh_token,
            "scope": payload.get("scope", tokens.scope),
            "token_type": payload.get("token_type", tokens.token_type),
            "expires_in": int(payload.get("expires_in", tokens.expires_in)),
        }

        return self._token_storage.save(
            google_id=tokens.google_id,
            display_name=tokens.display_name,
            email=tokens.email,
            payload=merged_payload,
        )

    async def _ensure_valid_tokens(self, tokens: StoredTokens) -> StoredTokens:
        if self._is_token_expired(tokens):
            return await self._refresh_access_token(tokens)
        return tokens

    async def _drive_request(
        self,
        tokens: StoredTokens,
        *,
        method: str,
        path: str,
        params: Optional[Dict[str, Any]] = None,
        json_body: Optional[Dict[str, Any]] = None,
    ) -> Tuple[Dict[str, Any], StoredTokens]:
        current_tokens = tokens

        for attempt in range(2):
            headers = {
                "Authorization": f"Bearer {current_tokens.access_token}",
                "Accept": "application/json",
            }

            async with httpx.AsyncClient(timeout=10.0, base_url=DRIVE_API_BASE) as client:
                response = await client.request(
                    method,
                    path,
                    params=params,
                    json=json_body,
                    headers=headers,
                )

            if response.status_code != 401:
                if response.is_error:
                    logger.error("Google Drive API %s %s failed: %s", method, path, response.text)
                    raise HTTPException(
                        status_code=502,
                        detail="Google Drive API 요청이 실패했습니다. 잠시 후 다시 시도해주세요.",
                    )

                data = response.json() if response.text else {}
                return data, current_tokens

            if attempt == 0:
                current_tokens = await self._refresh_access_token(current_tokens)
                continue

        raise HTTPException(status_code=401, detail="Google Drive 인증이 만료되었습니다. 다시 로그인해주세요.")

    async def _find_root_folder(
        self, tokens: StoredTokens, *, folder_name: str
    ) -> Tuple[Optional[Dict[str, Any]], StoredTokens]:
        escaped_name = folder_name.replace("'", "\\'")
        query = (
            f"name = '{escaped_name}' and "
            f"mimeType = '{DRIVE_FOLDER_MIME_TYPE}' and "
            "'root' in parents and trashed = false"
        )
        params = {
            "q": query,
            "fields": "files(id,name)",
            "pageSize": 1,
            "spaces": "drive",
        }

        data, updated_tokens = await self._drive_request(
            tokens,
            method="GET",
            path=DRIVE_FILES_ENDPOINT,
            params=params,
        )

        files = data.get("files")
        if isinstance(files, Sequence) and files:
            first = files[0]
            if isinstance(first, dict):
                return first, updated_tokens

        return None, updated_tokens

    async def _create_root_folder(
        self, tokens: StoredTokens, *, folder_name: str
    ) -> Tuple[Dict[str, Any], StoredTokens]:
        body = {
            "name": folder_name,
            "mimeType": DRIVE_FOLDER_MIME_TYPE,
            "parents": ["root"],
        }
        params = {"fields": "id,name"}

        data, updated_tokens = await self._drive_request(
            tokens,
            method="POST",
            path=DRIVE_FILES_ENDPOINT,
            params=params,
            json_body=body,
        )

        if not isinstance(data, dict) or "id" not in data:
            logger.error("Google Drive create folder response missing id: %s", data)
            raise HTTPException(status_code=502, detail="Google Drive 폴더를 생성하지 못했습니다. 다시 시도해주세요.")

        return data, updated_tokens

    async def _create_child_folder(
        self, tokens: StoredTokens, *, name: str, parent_id: str
    ) -> Tuple[Dict[str, Any], StoredTokens]:
        body = {
            "name": name,
            "mimeType": DRIVE_FOLDER_MIME_TYPE,
            "parents": [parent_id],
        }
        params = {"fields": "id,name,parents"}

        data, updated_tokens = await self._drive_request(
            tokens,
            method="POST",
            path=DRIVE_FILES_ENDPOINT,
            params=params,
            json_body=body,
        )

        if not isinstance(data, dict) or "id" not in data:
            logger.error("Google Drive create child folder response missing id: %s", data)
            raise HTTPException(status_code=502, detail="Google Drive 하위 폴더를 생성하지 못했습니다. 다시 시도해주세요.")

        return data, updated_tokens

    async def _upload_file_to_folder(
        self,
        tokens: StoredTokens,
        *,
        file_name: str,
        parent_id: str,
        content: bytes,
        content_type: Optional[str] = None,
    ) -> Tuple[Dict[str, Any], StoredTokens]:
        active_tokens = tokens

        for attempt in range(2):
            headers = {
                "Authorization": f"Bearer {active_tokens.access_token}",
            }
            metadata = {"name": file_name, "parents": [parent_id]}
            files = {
                "metadata": (
                    "metadata",
                    json.dumps(metadata),
                    "application/json; charset=UTF-8",
                ),
                "file": (
                    file_name,
                    content,
                    content_type or "application/pdf",
                ),
            }

            async with httpx.AsyncClient(timeout=30.0, base_url=DRIVE_UPLOAD_BASE) as client:
                response = await client.post(
                    f"{DRIVE_FILES_ENDPOINT}?uploadType=multipart&fields=id,name,parents",
                    headers=headers,
                    files=files,
                )

            if response.status_code == 401 and attempt == 0:
                active_tokens = await self._refresh_access_token(active_tokens)
                continue

            if response.is_error:
                logger.error("Google Drive file upload failed for %s: %s", file_name, response.text)
                raise HTTPException(
                    status_code=502,
                    detail="파일을 Google Drive에 업로드하지 못했습니다. 잠시 후 다시 시도해주세요.",
                )

            data = response.json()
            if not isinstance(data, dict) or "id" not in data:
                logger.error("Google Drive file upload response missing id: %s", data)
                raise HTTPException(status_code=502, detail="업로드한 파일의 ID를 확인하지 못했습니다. 다시 시도해주세요.")

            return data, active_tokens

        raise HTTPException(status_code=401, detail="Google Drive 인증이 만료되었습니다. 다시 로그인해주세요.")

    async def _download_file_content(
        self,
        tokens: StoredTokens,
        *,
        file_id: str,
        mime_type: Optional[str] = None,
    ) -> Tuple[bytes, StoredTokens]:
        active_tokens = tokens
        for attempt in range(2):
            headers = {
                "Authorization": f"Bearer {active_tokens.access_token}",
            }
            if mime_type == GOOGLE_SHEETS_MIME_TYPE:
                path = f"{DRIVE_FILES_ENDPOINT}/{file_id}/export"
                params = {"mimeType": XLSX_MIME_TYPE}
            else:
                path = f"{DRIVE_FILES_ENDPOINT}/{file_id}"
                params = {"alt": "media"}

            async with httpx.AsyncClient(timeout=30.0, base_url=DRIVE_API_BASE) as client:
                response = await client.get(
                    path,
                    params=params,
                    headers=headers,
                )

            if response.status_code == 401 and attempt == 0:
                active_tokens = await self._refresh_access_token(active_tokens)
                continue

            if response.is_error:
                logger.error("Google Drive file download failed for %s: %s", file_id, response.text)
                raise HTTPException(
                    status_code=502,
                    detail="Google Drive에서 파일을 다운로드하지 못했습니다. 잠시 후 다시 시도해주세요.",
                )

            return response.content, active_tokens

        raise HTTPException(status_code=401, detail="Google Drive 인증이 만료되었습니다. 다시 로그인해주세요.")

    async def _update_file_content(
        self,
        tokens: StoredTokens,
        *,
        file_id: str,
        file_name: str,
        content: bytes,
        content_type: str,
    ) -> Tuple[Dict[str, Any], StoredTokens]:
        active_tokens = tokens
        for attempt in range(2):
            headers = {
                "Authorization": f"Bearer {active_tokens.access_token}",
            }
            metadata = {"name": file_name}
            files = {
                "metadata": (
                    "metadata",
                    json.dumps(metadata),
                    "application/json; charset=UTF-8",
                ),
                "file": (
                    file_name,
                    content,
                    content_type,
                ),
            }

            async with httpx.AsyncClient(timeout=30.0, base_url=DRIVE_UPLOAD_BASE) as client:
                response = await client.patch(
                    f"{DRIVE_FILES_ENDPOINT}/{file_id}?uploadType=multipart&fields=id,name,modifiedTime",
                    headers=headers,
                    files=files,
                )

            if response.status_code == 401 and attempt == 0:
                active_tokens = await self._refresh_access_token(active_tokens)
                continue

            if response.is_error:
                logger.error("Google Drive file update failed for %s: %s", file_id, response.text)
                raise HTTPException(
                    status_code=502,
                    detail="Google Drive 파일을 업데이트하지 못했습니다. 잠시 후 다시 시도해주세요.",
                )

            data = response.json()
            if not isinstance(data, dict) or "id" not in data:
                logger.error("Google Drive file update response missing id: %s", data)
                raise HTTPException(status_code=502, detail="업데이트된 파일 정보를 확인하지 못했습니다. 다시 시도해주세요.")

            return data, active_tokens

        raise HTTPException(status_code=401, detail="Google Drive 인증이 만료되었습니다. 다시 로그인해주세요.")

    async def _resolve_menu_spreadsheet(
        self,
        *,
        project_id: str,
        menu_id: str,
        google_id: Optional[str],
        include_content: bool = False,
        file_id: Optional[str] = None,
    ) -> _ResolvedSpreadsheet:
        rule = _SPREADSHEET_RULES.get(menu_id)
        if not rule:
            raise HTTPException(status_code=404, detail="지원하지 않는 스프레드시트 메뉴입니다.")

        self._oauth_service.ensure_credentials()
        stored_tokens = self._load_tokens(google_id)
        active_tokens = await self._ensure_valid_tokens(stored_tokens)

        folder, active_tokens = await self._find_child_folder_by_name(
            active_tokens,
            parent_id=project_id,
            name=rule["folder_name"],
        )
        if folder is None or not folder.get("id"):
            raise HTTPException(status_code=404, detail=f"프로젝트에 '{rule['folder_name']}' 폴더를 찾을 수 없습니다.")

        folder_id = str(folder["id"])
        file_entry: Optional[Dict[str, Any]] = None
        if file_id:
            file_entry, active_tokens = await self._get_file_metadata(
                active_tokens,
                file_id=file_id,
            )
            if file_entry is None or not file_entry.get("id"):
                raise HTTPException(status_code=404, detail=f"프로젝트에 '{rule['file_suffix']}' 파일을 찾을 수 없습니다.")

            parents = file_entry.get("parents")
            if isinstance(parents, Sequence) and parents:
                parent_ids = {
                    parent.decode("utf-8") if isinstance(parent, bytes) else str(parent)
                    for parent in parents
                    if isinstance(parent, (str, bytes))
                }
                if folder_id not in parent_ids:
                    logger.warning(
                        "Drive file is outside expected folder",
                        extra={
                            "project_id": project_id,
                            "menu_id": menu_id,
                            "expected_folder_id": folder_id,
                            "file_parents": list(parent_ids),
                            "file_id": file_id,
                        },
                    )
        else:
            file_entry, active_tokens = await self._find_file_by_suffix(
                active_tokens,
                parent_id=folder_id,
                suffix=rule["file_suffix"],
                mime_type=XLSX_MIME_TYPE,
            )
            if file_entry is None or not file_entry.get("id"):
                raise HTTPException(status_code=404, detail=f"프로젝트에 '{rule['file_suffix']}' 파일을 찾을 수 없습니다.")

        file_id = str(file_entry["id"])
        file_name = str(file_entry.get("name", rule["file_suffix"]))
        mime_type = file_entry.get("mimeType")
        normalized_mime = mime_type if isinstance(mime_type, str) else None
        modified_time = (
            str(file_entry.get("modifiedTime"))
            if isinstance(file_entry.get("modifiedTime"), str)
            else None
        )

        content: Optional[bytes] = None
        if include_content:
            content, active_tokens = await self._download_file_content(
                active_tokens,
                file_id=file_id,
                mime_type=normalized_mime,
            )

        return _ResolvedSpreadsheet(
            rule=rule,
            tokens=active_tokens,
            folder_id=folder_id,
            file_id=file_id,
            file_name=file_name,
            mime_type=normalized_mime,
            modified_time=modified_time,
            content=content,
        )

    async def apply_csv_to_spreadsheet(
        self,
        *,
        project_id: str,
        menu_id: str,
        csv_text: str,
        google_id: Optional[str],
        project_overview: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:
        rule = _SPREADSHEET_RULES.get(menu_id)
        if not rule:
            return None

        resolved = await self._resolve_menu_spreadsheet(
            project_id=project_id,
            menu_id=menu_id,
            google_id=google_id,
            include_content=True,
        )

        workbook_bytes = resolved.content
        if workbook_bytes is None:
            raise HTTPException(status_code=500, detail="스프레드시트 내용을 불러오지 못했습니다. 다시 시도해 주세요.")

        overview_value: Optional[str] = None
        try:
            populate = resolved.rule["populate"]
            if menu_id == "feature-list":
                overview_value = (
                    str(project_overview or "") if project_overview is not None else None
                )
                updated_bytes = populate(workbook_bytes, csv_text, overview_value)
            else:
                updated_bytes = populate(workbook_bytes, csv_text)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        except Exception as exc:  # pragma: no cover - 안전망
            logger.exception(
                "Failed to populate spreadsheet for project", extra={"project_id": project_id, "menu_id": menu_id}
            )
            raise HTTPException(status_code=500, detail="엑셀 템플릿을 업데이트하지 못했습니다. 다시 시도해주세요.") from exc

        update_info, _ = await self._update_file_content(
            resolved.tokens,
            file_id=resolved.file_id,
            file_name=resolved.file_name,
            content=updated_bytes,
            content_type=XLSX_MIME_TYPE,
        )
        logger.info(
            "Populated project spreadsheet",
            extra={"project_id": project_id, "menu_id": menu_id, "file_id": resolved.file_id},
        )
        response: Dict[str, Any] = {
            "fileId": resolved.file_id,
            "fileName": resolved.file_name,
            "modifiedTime": update_info.get("modifiedTime") if isinstance(update_info, dict) else None,
        }
        if menu_id == "feature-list" and overview_value is not None:
            response["projectOverview"] = overview_value
        return response

    async def get_feature_list_rows(
        self,
        *,
        project_id: str,
        google_id: Optional[str],
        file_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        resolved = await self._resolve_menu_spreadsheet(
            project_id=project_id,
            menu_id="feature-list",
            google_id=google_id,
            include_content=True,
            file_id=file_id,
        )

        workbook_bytes = resolved.content
        if workbook_bytes is None:
            raise HTTPException(status_code=500, detail="기능리스트 파일을 불러오지 못했습니다. 다시 시도해 주세요.")

        _, project_overview = extract_feature_list_overview(workbook_bytes)

        buffer = io.BytesIO(workbook_bytes)
        try:
            workbook = load_workbook(buffer, data_only=True)
        except Exception as exc:  # pragma: no cover - 안전망
            raise HTTPException(status_code=500, detail="엑셀 파일을 읽는 중 오류가 발생했습니다.") from exc

        headers = list(FEATURE_LIST_EXPECTED_HEADERS)
        extracted_rows: List[Dict[str, str]] = []
        sheet_title = ""
        start_row = _FEATURE_LIST_START_ROW
        header_row_values: Optional[Sequence[Any]] = None
        column_map: Dict[str, int] = {}
        try:
            sheet = workbook.active
            selected_title = sheet.title
            for candidate in _FEATURE_LIST_SHEET_CANDIDATES:
                matched = False
                for title in workbook.sheetnames:
                    if _drive_name_matches(title, candidate):
                        try:
                            sheet = workbook[title]
                            selected_title = sheet.title
                            matched = True
                            break
                        except KeyError:
                            continue
                if matched:
                    break

            sheet_title = selected_title or ""
            max_col = max(len(headers), sheet.max_column or len(headers))
            header_row_index: Optional[int] = None
            first_data_row_index: Optional[int] = None
            for idx, row in enumerate(
                sheet.iter_rows(min_row=1, max_col=max_col, values_only=True),
                start=1,
            ):
                row_values: Sequence[Any] = row if isinstance(row, Sequence) else tuple()

                has_values = False
                for col_idx in range(len(headers)):
                    cell_value = row_values[col_idx] if col_idx < len(row_values) else None
                    if cell_value is None:
                        continue
                    if str(cell_value).strip():
                        has_values = True
                        break

                header_match = _looks_like_header_row(row_values, headers)

                if has_values and not header_match and first_data_row_index is None:
                    first_data_row_index = idx

                if header_match:
                    header_row_index = idx
                    header_row_values = row_values
                    break

                if idx >= _FEATURE_LIST_START_ROW * 2 and first_data_row_index is not None:
                    break

            if header_row_index is not None:
                start_row = header_row_index + 1
            elif first_data_row_index is not None:
                start_row = max(1, first_data_row_index)

            if header_row_values:
                display_headers = list(headers)
                for idx, value in enumerate(header_row_values):
                    if value is None:
                        continue
                    matched = match_feature_list_header(str(value))
                    if matched and matched not in column_map:
                        column_map[matched] = idx
                        try:
                            header_index = headers.index(matched)
                        except ValueError:
                            header_index = None
                        if header_index is not None:
                            display_headers[header_index] = str(value).strip()

                headers = display_headers

            for default_idx, name in enumerate(FEATURE_LIST_EXPECTED_HEADERS):
                column_map.setdefault(name, default_idx)

            for row in sheet.iter_rows(
                min_row=max(1, start_row),
                max_col=max_col,
                values_only=True,
            ):
                row_values: Sequence[Any] = row if isinstance(row, Sequence) else tuple()

                if _looks_like_header_row(row_values, headers):
                    continue

                row_data: Dict[str, str] = {}
                has_values = False
                for header_name in headers:
                    column_index = column_map.get(header_name)
                    cell_value = (
                        row_values[column_index]
                        if column_index is not None and column_index < len(row_values)
                        else None
                    )
                    text = "" if cell_value is None else str(cell_value).strip()
                    if text:
                        has_values = True
                    row_data[header_name] = text

                if not has_values:
                    continue

                description = row_data.get("기능 설명", "")
                if not description:
                    description = row_data.get("기능 개요", "")

                extracted_rows.append(
                    {
                        "majorCategory": row_data.get("대분류", ""),
                        "middleCategory": row_data.get("중분류", ""),
                        "minorCategory": row_data.get("소분류", ""),
                        "featureDescription": description,
                    }
                )
        finally:
            workbook.close()

        if not sheet_title:
            sheet_title = "기능리스트"

        return {
            "fileId": resolved.file_id,
            "fileName": resolved.file_name,
            "sheetName": sheet_title,
            "startRow": start_row,
            "headers": headers,
            "rows": extracted_rows,
            "modifiedTime": resolved.modified_time,
            "projectOverview": project_overview,
        }

    async def update_feature_list_rows(
        self,
        *,
        project_id: str,
        rows: Sequence[Dict[str, str]],
        project_overview: str = "",
        google_id: Optional[str],
        file_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        resolved = await self._resolve_menu_spreadsheet(
            project_id=project_id,
            menu_id="feature-list",
            google_id=google_id,
            include_content=True,
            file_id=file_id,
        )

        workbook_bytes = resolved.content
        if workbook_bytes is None:
            raise HTTPException(status_code=500, detail="기능리스트 파일을 불러오지 못했습니다. 다시 시도해 주세요.")

        output = io.StringIO()
        writer = csv.DictWriter(
            output,
            fieldnames=list(FEATURE_LIST_EXPECTED_HEADERS),
            lineterminator="\n",
        )
        writer.writeheader()

        has_overview_column = "기능 개요" in FEATURE_LIST_EXPECTED_HEADERS

        for row in rows:
            major = str(row.get("majorCategory", "") or "").strip()
            middle = str(row.get("middleCategory", "") or "").strip()
            minor = str(row.get("minorCategory", "") or "").strip()
            description = str(row.get("featureDescription", "") or "").strip()

            if not any([major, middle, minor, description]):
                continue

            entry = {
                "대분류": major,
                "중분류": middle,
                "소분류": minor,
                "기능 설명": description,
            }
            if has_overview_column:
                entry["기능 개요"] = ""

            writer.writerow(entry)

        csv_text = output.getvalue()

        try:
            updated_bytes = resolved.rule["populate"](
                workbook_bytes,
                csv_text,
                project_overview,
            )
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        except Exception as exc:  # pragma: no cover - 안전망
            logger.exception("Failed to update feature list spreadsheet", extra={"project_id": project_id})
            raise HTTPException(status_code=500, detail="기능리스트를 업데이트하지 못했습니다. 다시 시도해 주세요.") from exc

        update_info, _ = await self._update_file_content(
            resolved.tokens,
            file_id=resolved.file_id,
            file_name=resolved.file_name,
            content=updated_bytes,
            content_type=XLSX_MIME_TYPE,
        )

        return {
            "fileId": resolved.file_id,
            "fileName": resolved.file_name,
            "modifiedTime": update_info.get("modifiedTime") if isinstance(update_info, dict) else None,
            "projectOverview": project_overview,
        }

    async def download_feature_list_workbook(
        self,
        *,
        project_id: str,
        google_id: Optional[str],
        file_id: Optional[str] = None,
    ) -> Tuple[str, bytes]:
        resolved = await self._resolve_menu_spreadsheet(
            project_id=project_id,
            menu_id="feature-list",
            google_id=google_id,
            include_content=True,
            file_id=file_id,
        )

        workbook_bytes = resolved.content
        if workbook_bytes is None:
            raise HTTPException(status_code=500, detail="기능리스트 파일을 불러오지 못했습니다. 다시 시도해 주세요.")

        return resolved.file_name, workbook_bytes

    async def get_project_exam_number(
        self,
        *,
        project_id: str,
        google_id: Optional[str],
    ) -> str:
        """
        Retrieve the exam number (e.g. GS-B-12-3456) from the Drive project folder name.
        """
        self._oauth_service.ensure_credentials()
        stored_tokens = self._load_tokens(google_id)
        active_tokens = await self._ensure_valid_tokens(stored_tokens)

        params = {"fields": "id,name"}
        data, _ = await self._drive_request(
            active_tokens,
            method="GET",
            path=f"{DRIVE_FILES_ENDPOINT}/{project_id}",
            params=params,
        )

        name = data.get("name")
        if not isinstance(name, str) or not name.strip():
            raise HTTPException(status_code=404, detail="프로젝트 폴더를 찾을 수 없습니다.")

        match = EXAM_NUMBER_PATTERN.search(name)
        if not match:
            raise HTTPException(status_code=404, detail="프로젝트 이름에서 시험신청 번호를 찾을 수 없습니다.")

        return match.group(0)

    async def _ensure_shared_criteria_file(
        self,
        tokens: StoredTokens,
        *,
        parent_id: str,
        preferred_names: Optional[Sequence[str]] = None,
    ) -> Tuple[Dict[str, Any], StoredTokens, bool]:
        normalized_candidates = set(_SHARED_CRITERIA_NORMALIZED_NAMES)
        upload_name = _PREFERRED_SHARED_CRITERIA_FILE_NAME
        if preferred_names:
            normalized_candidates.update(
                _normalize_shared_criteria_name(name)
                for name in preferred_names
                if isinstance(name, str) and name.strip()
            )
            first_valid = next(
                (name.strip() for name in preferred_names if isinstance(name, str) and name.strip()),
                None,
            )
            if first_valid:
                upload_name = first_valid

        files, active_tokens = await self._list_child_files(tokens, parent_id=parent_id)
        for entry in files:
            if not isinstance(entry, dict):
                continue
            name = entry.get("name")
            if not isinstance(name, str):
                continue
            normalized_name = _normalize_shared_criteria_name(name)
            if normalized_name not in normalized_candidates:
                continue
            mime_type = entry.get("mimeType")
            if isinstance(mime_type, str) and mime_type not in {XLSX_MIME_TYPE, GOOGLE_SHEETS_MIME_TYPE}:
                logger.warning(
                    "Ignoring shared criteria candidate with unsupported mime type: %s (%s)",
                    name,
                    mime_type,
                )
                continue
            normalized_entry = dict(entry)
            normalized_entry["mimeType"] = mime_type if isinstance(mime_type, str) else None
            return normalized_entry, active_tokens, False

        content = GoogleDriveService._load_shared_criteria_template_bytes()
        uploaded_entry, updated_tokens = await self._upload_file_to_folder(
            active_tokens,
            file_name=upload_name,
            parent_id=parent_id,
            content=content,
            content_type=XLSX_MIME_TYPE,
        )
        uploaded_entry = dict(uploaded_entry)
        uploaded_entry.setdefault("name", upload_name)
        uploaded_entry["mimeType"] = XLSX_MIME_TYPE
        logger.info(
            "Uploaded shared criteria template to gs folder: %s",
            uploaded_entry.get("name"),
        )
        return uploaded_entry, updated_tokens, True

    async def download_shared_security_criteria(
        self,
        *,
        google_id: Optional[str],
        file_name: str,
    ) -> bytes:
        self._oauth_service.ensure_credentials()
        stored_tokens = self._load_tokens(google_id)
        active_tokens = await self._ensure_valid_tokens(stored_tokens)

        folder, active_tokens = await self._find_root_folder(active_tokens, folder_name="gs")
        if folder is None:
            folder, active_tokens = await self._create_root_folder(active_tokens, folder_name="gs")
        gs_folder_id = str(folder["id"])

        file_entry, active_tokens, _ = await self._ensure_shared_criteria_file(
            active_tokens,
            parent_id=gs_folder_id,
            preferred_names=(file_name,),
        )

        file_id = file_entry.get("id")
        if not isinstance(file_id, str):
            logger.error("Shared criteria entry missing id: %s", file_entry)
            raise HTTPException(status_code=502, detail="결함 판단 기준표 ID를 확인할 수 없습니다.")

        content, _ = await self._download_file_content(
            active_tokens,
            file_id=file_id,
            mime_type=file_entry.get("mimeType"),
        )
        return content

    async def _list_child_folders(
        self,
        tokens: StoredTokens,
        *,
        parent_id: str,
    ) -> Tuple[Sequence[Dict[str, Any]], StoredTokens]:
        query = (
            f"'{parent_id}' in parents and "
            f"mimeType = '{DRIVE_FOLDER_MIME_TYPE}' and trashed = false"
        )
        params = {
            "q": query,
            "fields": "files(id,name,createdTime,modifiedTime)",
            "orderBy": "name_natural",
            "spaces": "drive",
            "pageSize": 100,
            "supportsAllDrives": "true",
            "includeItemsFromAllDrives": "true",
        }

        data, updated_tokens = await self._drive_request(
            tokens,
            method="GET",
            path=DRIVE_FILES_ENDPOINT,
            params=params,
        )

        files = data.get("files")
        if isinstance(files, Sequence):
            return files, updated_tokens

        return [], updated_tokens

    async def _list_child_files(
        self,
        tokens: StoredTokens,
        *,
        parent_id: str,
        mime_type: Optional[str] = None,
    ) -> Tuple[Sequence[Dict[str, Any]], StoredTokens]:
        clauses = [f"'{parent_id}' in parents", "trashed = false"]
        if mime_type:
            clauses.append(f"mimeType = '{mime_type}'")
        params = {
            "q": " and ".join(clauses),
            "fields": "files(id,name,mimeType,modifiedTime)",
            "orderBy": "name_natural",
            "spaces": "drive",
            "pageSize": 100,
            "supportsAllDrives": "true",
            "includeItemsFromAllDrives": "true",
        }

        data, updated_tokens = await self._drive_request(
            tokens,
            method="GET",
            path=DRIVE_FILES_ENDPOINT,
            params=params,
        )

        files = data.get("files")
        if isinstance(files, Sequence):
            return files, updated_tokens

        return [], updated_tokens

    async def _find_child_folder_by_name(
        self,
        tokens: StoredTokens,
        *,
        parent_id: str,
        name: str,
    ) -> Tuple[Optional[Dict[str, Any]], StoredTokens]:
        folders, updated_tokens = await self._list_child_folders(tokens, parent_id=parent_id)
        target_variants = set(_drive_name_variants(name))
        for folder in folders:
            if not isinstance(folder, dict):
                continue
            folder_name = folder.get("name")
            if not isinstance(folder_name, str):
                continue
            if folder_name == name:
                return folder, updated_tokens
            if target_variants and set(_drive_name_variants(folder_name)) & target_variants:
                return folder, updated_tokens
        return None, updated_tokens

    async def _find_file_by_suffix(
        self,
        tokens: StoredTokens,
        *,
        parent_id: str,
        suffix: str,
        mime_type: Optional[str] = None,
    ) -> Tuple[Optional[Dict[str, Any]], StoredTokens]:
        search_mime_types: Sequence[Optional[str]]
        if mime_type:
            search_mime_types = (mime_type, None)
        else:
            search_mime_types = (None,)

        updated_tokens = tokens
        for candidate_mime in search_mime_types:
            files, updated_tokens = await self._list_child_files(
                updated_tokens,
                parent_id=parent_id,
                mime_type=candidate_mime,
            )
            for entry in files:
                if not isinstance(entry, dict):
                    continue
                name = entry.get("name")
                if isinstance(name, str):
                    if name.endswith(suffix.strip()) or _drive_suffix_matches(name, suffix):
                        return entry, updated_tokens
        return None, updated_tokens

    async def _get_file_metadata(
        self,
        tokens: StoredTokens,
        *,
        file_id: str,
    ) -> Tuple[Optional[Dict[str, Any]], StoredTokens]:
        active_tokens = tokens
        params = {
            "fields": "id,name,mimeType,modifiedTime,parents",
            "supportsAllDrives": "true",
        }
        for attempt in range(2):
            headers = {
                "Authorization": f"Bearer {active_tokens.access_token}",
                "Accept": "application/json",
            }

            async with httpx.AsyncClient(timeout=10.0, base_url=DRIVE_API_BASE) as client:
                response = await client.get(
                    f"{DRIVE_FILES_ENDPOINT}/{file_id}",
                    params=params,
                    headers=headers,
                )

            if response.status_code == 401 and attempt == 0:
                active_tokens = await self._refresh_access_token(active_tokens)
                continue

            if response.status_code == 404:
                return None, active_tokens

            if response.is_error:
                logger.error("Google Drive metadata fetch failed for %s: %s", file_id, response.text)
                raise HTTPException(
                    status_code=502,
                    detail="Google Drive에서 파일 정보를 불러오지 못했습니다. 잠시 후 다시 시도해주세요.",
                )

            data = response.json() if response.text else {}
            if not isinstance(data, dict):
                logger.error("Google Drive metadata response malformed for %s: %s", file_id, data)
                raise HTTPException(
                    status_code=502,
                    detail="Google Drive 파일 정보를 확인하지 못했습니다. 다시 시도해주세요.",
                )

            return data, active_tokens

        raise HTTPException(status_code=401, detail="Google Drive 인증이 만료되었습니다. 다시 로그인해주세요.")

    async def _find_file_by_name(
        self,
        tokens: StoredTokens,
        *,
        parent_id: str,
        name: str,
        mime_type: Optional[str] = None,
    ) -> Tuple[Optional[Dict[str, Any]], StoredTokens]:
        files, updated_tokens = await self._list_child_files(
            tokens, parent_id=parent_id, mime_type=mime_type
        )
        normalized_name = name.strip()
        for entry in files:
            if not isinstance(entry, dict):
                continue
            file_name = entry.get("name")
            if isinstance(file_name, str) and file_name.strip() == normalized_name:
                return entry, updated_tokens
        return None, updated_tokens

    async def ensure_drive_setup(self, google_id: Optional[str]) -> Dict[str, Any]:
        self._oauth_service.ensure_credentials()
        stored_tokens = self._load_tokens(google_id)
        active_tokens = await self._ensure_valid_tokens(stored_tokens)

        folder, active_tokens = await self._find_root_folder(active_tokens, folder_name="gs")
        folder_created = False

        if folder is None:
            folder, active_tokens = await self._create_root_folder(
                active_tokens, folder_name="gs"
            )
            folder_created = True

        gs_folder_id = str(folder["id"])

        criteria_sheet, active_tokens, criteria_created = await self._ensure_shared_criteria_file(
            active_tokens,
            parent_id=gs_folder_id,
        )

        projects, active_tokens = await self._list_child_folders(
            active_tokens, parent_id=str(folder["id"])
        )

        normalized_projects = []
        for item in projects:
            if not isinstance(item, dict):
                continue
            project_id = item.get("id")
            name = item.get("name")
            if not isinstance(project_id, str) or not isinstance(name, str):
                continue
            normalized_projects.append(
                {
                    "id": project_id,
                    "name": name,
                    "createdTime": item.get("createdTime"),
                    "modifiedTime": item.get("modifiedTime"),
                }
            )

        return {
            "folderCreated": folder_created,
            "folderId": folder["id"],
            "folderName": folder.get("name", "gs"),
            "criteria": {
                "created": criteria_created,
                "fileId": criteria_sheet.get("id"),
                "fileName": criteria_sheet.get("name"),
                "mimeType": criteria_sheet.get("mimeType"),
            },
            "projects": normalized_projects,
            "account": {
                "googleId": active_tokens.google_id,
                "displayName": active_tokens.display_name,
                "email": active_tokens.email,
            },
        }

    async def create_project(
        self,
        *,
        folder_id: Optional[str],
        files: List[UploadFile],
        google_id: Optional[str],
    ) -> Dict[str, Any]:
        self._oauth_service.ensure_credentials()
        stored_tokens = self._load_tokens(google_id)
        active_tokens = await self._ensure_valid_tokens(stored_tokens)

        parent_folder_id = folder_id
        if not parent_folder_id:
            folder, active_tokens = await self._find_root_folder(active_tokens, folder_name="gs")
            if folder is None:
                folder, active_tokens = await self._create_root_folder(active_tokens, folder_name="gs")
            parent_folder_id = str(folder["id"])

        agreement_file = files[0]
        if not agreement_file.filename or not agreement_file.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=422, detail="시험 합의서는 DOCX 파일이어야 합니다.")

        agreement_bytes = await agreement_file.read()
        metadata = self._extract_project_metadata(agreement_bytes)
        project_name = self._build_project_folder_name(metadata)
        if not project_name:
            raise HTTPException(status_code=422, detail="생성할 프로젝트 이름을 결정할 수 없습니다.")

        siblings, active_tokens = await self._list_child_folders(active_tokens, parent_id=parent_folder_id)
        existing_names = {
            str(item.get("name"))
            for item in siblings
            if isinstance(item, dict) and isinstance(item.get("name"), str)
        }

        unique_name = project_name
        suffix = 1
        while unique_name in existing_names:
            suffix += 1
            unique_name = f"{project_name} ({suffix})"

        project_folder, active_tokens = await self._create_child_folder(
            active_tokens,
            name=unique_name,
            parent_id=parent_folder_id,
        )
        project_id = str(project_folder["id"])

        active_tokens = await self._copy_template_to_drive(
            active_tokens,
            parent_id=project_id,
            exam_number=metadata["exam_number"],
        )

        uploaded_files: List[Dict[str, Any]] = []

        agreement_name = agreement_file.filename or "시험 합의서.docx"
        agreement_name = self._replace_placeholders(agreement_name, metadata["exam_number"])
        file_info, active_tokens = await self._upload_file_to_folder(
            active_tokens,
            file_name=agreement_name,
            parent_id=project_id,
            content=agreement_bytes,
            content_type=agreement_file.content_type
            or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        uploaded_files.append(
            {
                "id": file_info.get("id"),
                "name": file_info.get("name", agreement_name),
                "size": len(agreement_bytes),
                "contentType": agreement_file.content_type
                or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }
        )
        await agreement_file.close()

        for upload in files[1:]:
            filename = upload.filename or "업로드된 파일.docx"
            content = await upload.read()
            file_info, active_tokens = await self._upload_file_to_folder(
                active_tokens,
                file_name=filename,
                parent_id=project_id,
                content=content,
                content_type=upload.content_type,
            )
            uploaded_files.append(
                {
                    "id": file_info.get("id"),
                    "name": file_info.get("name", filename),
                    "size": len(content),
                    "contentType": upload.content_type or "application/octet-stream",
                }
            )
            await upload.close()

        logger.info(
            "Created Drive project '%s' (%s) with metadata %s",
            unique_name,
            project_id,
            metadata,
        )

        return {
            "message": "새 프로젝트 폴더를 생성했습니다.",
            "project": {
                "id": project_id,
                "name": project_folder.get("name", unique_name),
                "parentId": parent_folder_id,
                "metadata": {
                    "examNumber": metadata["exam_number"],
                    "companyName": metadata["company_name"],
                    "productName": metadata["product_name"],
                },
            },
            "uploadedFiles": uploaded_files,
        }
